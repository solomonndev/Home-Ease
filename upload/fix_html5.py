"""Fix the remaining HTML5 paragraph that's split across runs."""
from docx import Document

INPUT = "/home/z/my-project/upload/new mariam.docx"
OUTPUT = "/home/z/my-project/upload/HomeEase_Project_Writeup.docx"

doc = Document(INPUT)

# Find the paragraph containing "HTML5" and force-replace it
fixed = 0
for para in doc.paragraphs:
    full = "".join(r.text for r in para.runs)
    if "HTML5" in full:
        print(f"FOUND HTML5 paragraph (length={len(full)}):")
        print(f"  '{full[:200]}...'")
        print(f"  '{full[-100:]}'")
        
        # The full replacement
        new_text = (
            "The system implementation follows an agile development methodology [31], "
            "enabling incremental development and continuous feedback. The front end "
            "of the virtual space is developed using Next.js 16 with TypeScript and Tailwind CSS 4 "
            "for type-safe, responsive, and dynamic user interfaces. The backend system "
            "is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, "
            "which manages business logic, service scheduling, user authentication, real-time "
            "messaging via Socket.io, and communication between system components. Data is stored "
            "and managed using SQLite as the relational database engine, with Prisma ORM providing "
            "type-safe schema management and an intuitive query API for data access."
        )
        
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(new_text)
        fixed += 1
        print("  ✅ REPLACED!")

# Also check table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full = "".join(r.text for r in para.runs)
                if "HTML5" in full:
                    new_text = full.replace("HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces",
                                           "Next.js 16 + TypeScript + Tailwind CSS 4 + shadcn/ui for responsive and dynamic user interfaces")
                    if para.runs:
                        para.runs[0].text = new_text
                        for r in para.runs[1:]:
                            r.text = ""
                    fixed += 1
                    print("  ✅ Fixed HTML5 in table!")

print(f"\nTotal HTML5 fixes: {fixed}")

doc.save(OUTPUT)
print(f"Saved to: {OUTPUT}")

# Verify
import subprocess
r = subprocess.run(["pandoc", OUTPUT, "-t", "plain"], capture_output=True, text=True)
t = r.stdout
for term in ["HTML5, CSS3", "Node.js with Express", "Next.js 16", "TypeScript", "Tailwind CSS 4", "Bun", "Prisma ORM", "SQLite", "Socket.io", "JWT"]:
    print(f"  {'✅' if term in t else '❌'} {term}")
