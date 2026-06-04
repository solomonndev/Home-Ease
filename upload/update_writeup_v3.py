"""
Final robust script - force full-paragraph text replacement for reliability.
"""
from docx import Document

INPUT_FILE = "/home/z/my-project/upload/new mariam.docx"
OUTPUT_FILE = "/home/z/my-project/upload/HomeEase_Project_Writeup.docx"

doc = Document(INPUT_FILE)

def force_replace(doc, old_text, new_text):
    """Aggressively replace text across all paragraphs and table cells.
    Rebuilds the paragraph with combined text when match is found."""
    count = 0
    
    targets = []  # (para, old, new) list
    
    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs)
        if old_text in full:
            targets.append((para, old_text, new_text))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = "".join(r.text for r in para.runs)
                    if old_text in full:
                        targets.append((para, old_text, new_text))
    
    for para, old, new in targets:
        full = "".join(r.text for r in para.runs)
        new_full = full.replace(old, new)
        
        # Put all text on first run, clear rest
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(new_full)
        count += 1
    
    return count

# ============================================================
# ALL TECH STACK REPLACEMENTS (longest match first to avoid conflicts)
# ============================================================

stack_replacements = [
    # Section 1.E - The big methodology paragraph
    (
        "The front end of the virtual space is developed using HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces. The backend system is implemented using Node.js with Express which manage business logic, service scheduling, user authentication, and communication between system components. Data is stored and managed using a relational database such as MySQL for scalability and data structure requirements.",
        "The front end of the virtual space is developed using Next.js 16 with TypeScript and Tailwind CSS 4 for type-safe, responsive, and dynamic user interfaces. The backend system is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, which manages business logic, service scheduling, user authentication, real-time messaging via Socket.io, and communication between system components. Data is stored and managed using SQLite as the relational database engine, with Prisma ORM providing type-safe schema management and an intuitive query API for data access."
    ),
    (
        "Firebase is used for secure user authentication and role-based access control, while Stripe or Paystack APIs are employed for secure payment processing.",
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication and role-based access control, while the Paystack payment gateway API is employed for secure payment processing integrated with a virtual wallet and escrow system."
    ),
    (
        "Cloud services such as Google Cloud, or Hostinger support hosting, storage, and scalability of the virtual space.",
        "Vercel, a cloud platform optimized for Next.js applications, supports hosting, serverless function execution, and automatic scalability of the virtual space."
    ),
    (
        "software testing tools such as Jest, Selenium, or Postman to verify functionality, performance, and security.",
        "software testing tools such as Jest for unit testing and Postman for API endpoint verification to verify functionality, performance, and security."
    ),
    # Section 3.4 intro
    (
        "Commonly used DBMS solutions such as MySQL or MongoDB are suitable due to their reliability and scalability.",
        "SQLite is used as the relational database engine, managed through Prisma ORM which provides type-safe database access and schema management, ensuring reliability and data integrity."
    ),
]

print("=== Applying Tech Stack Replacements ===")
for old, new in stack_replacements:
    c = force_replace(doc, old, new)
    if c > 0:
        print(f"  ✅ Replaced ({c}): '{old[:60]}...'")
    else:
        print(f"  ❌ NOT FOUND: '{old[:60]}...'")

# ============================================================
# TABLE-SPECIFIC FIXES
# ============================================================
print("\n=== Fixing Software Requirements Table ===")
table_specific = [
    ("MongoDB(NOSQL)", "SQLite with Prisma ORM"),
    ("MongoDB (NOSQL)", "SQLite with Prisma ORM"),
    ("React native +", "Next.js 16 + TypeScript +"),
    ("Node.js with", "Next.js 16 (App Router) +"),
    ("Paystack or", "Paystack Payment"),
    ("Flutterwave", "Gateway"),
]

for old, new in table_specific:
    c = force_replace(doc, old, new)
    if c > 0:
        print(f"  ✅ Table fix ({c}): '{old}' → '{new}'")
    else:
        print(f"  ⚠️ Not found: '{old}'")

# ============================================================
# INLINE REFERENCES  
# ============================================================
print("\n=== Applying Inline References ===")

refs = [
    # Ch1 Background
    ("The rapid advancement of digital technologies has transformed how services are delivered across various industries. Domestic",
     "The rapid advancement of digital technologies has transformed how services are delivered across various industries [1]. Domestic"),
    ("Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping. Although",
     "Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping [2]. Although"),
    ("The concept of a virtual space for domestic service involves creating an online platform or digital environment",
     "The concept of a virtual space for domestic service involves creating an online platform or digital environment [3]"),
    ("Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance, developing",
     "Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance [4], developing"),
    
    # Ch1 Problem Statement
    ("The provision of domestic services is largely managed through informal and manual processes, resulting",
     "The provision of domestic services is largely managed through informal and manual processes [5, 6], resulting"),

    # Ch1 Lit table 
    ("Khatri & Gupta (2020)", "Khatri and Gupta [5]"),
    ("Adeyemi & Fatile\n                (2021)", "Adeyemi and Fatile\n                [6]"),
    ("Rana et al.\n   (2019)", "Rana et al.\n   [7]"),
    ("Indravasan\n   et al.\n                (2018)", "Indravasan\n   et al.\n                [8]"),
    ("Chen et al. (2021)", "Chen et al. [9]"),

    # Ch2 - TAM
    ("perceived ease of use (Davis, 1989). In the context",
     "perceived ease of use [10]. In the context"),
    
    # Ch2 - Platform Economy
    ("The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries",
     "The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries [11]"),
    ("Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms. Dimensions",
     "Additionally, Service Quality Theory [12] explains how users evaluate services delivered through digital platforms. Dimensions"),

    # Ch2 - Virtual space / payments
    ("When integrated with Paystack digital payment system, the",
     "When integrated with a digital payment system such as Paystack [14], the"),
    ("Digital payment systems have achieved widespread adoption due to their",
     "Digital payment systems have achieved widespread adoption due to their [14]"),

    # Ch2 - Web 3.0
    ("Web 3.0 refers to the next stage in the evolution of the internet, designed to move",
     "Web 3.0 refers to the next stage in the evolution of the internet [15], designed to move"),
    ("The heart of Web 3.0 is blockchain technology. A blockchain is a distributed ledger",
     "The heart of Web 3.0 is blockchain technology. A blockchain is a distributed ledger [16]"),

    # Ch2 - Database
    ("A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated. Databases are",
     "A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated [17]. Databases are"),
    ("ACID properties are fundamental principles that ensure the reliability and correctness of database transactions. A transaction refers",
     "ACID properties are fundamental principles that ensure the reliability and correctness of database transactions [20]. A transaction refers"),

    # Ch2 - Data modeling
    ("Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system. It defines",
     "Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system [24]. It defines"),
    ("Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures. It is widely",
     "Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures [25]. It is widely"),
    ("Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency. It involves",
     "Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency [26]. It involves"),

    # Ch2 - SDLC
    ("A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems. It provides",
     "A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems [27]. It provides"),
    ("System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner. It provides a step-by-step approach",
     "System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner [28]. It provides a step-by-step approach"),
    ("The waterfall model is a sequential, plan driven-process where you must",
     "The waterfall model is a sequential, plan-driven process [29] where you must"),
    ("The V model (Verification and Validation model) is an extension of the waterfall model. All the requirements are gathered at the start",
     "The V model (Verification and Validation model) is an extension of the waterfall model [30]. All the requirements are gathered at the start"),
    ("The agile process model encourages continuous iterations of development and testing.",
     "The agile process model [31] encourages continuous iterations of development and testing."),

    # Ch2 - Hosting
    ("Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure. These infrastructures",
     "Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure [32]. These infrastructures"),
    ("Platform as a Service (PaaS) provides a complete hosting environment that includes operating systems, development frameworks, databases, and deployment tools.",
     "Platform as a Service (PaaS) [33] provides a complete hosting environment that includes operating systems, development frameworks, databases, and deployment tools."),

    # Ch2 - Related Works
    ("Aishwaryalakshmi et al. (2024), focuses on", "Aishwaryalakshmi et al. [34], focuses on"),
    ("Pais and Zanoni (2024), Their study adopted", "Pais and Zanoni [35]. Their study adopted"),
    ("Meyanban et al. (2024), This study explored", "Meyanban et al. [36]. This study explored"),
    ("Orth and Baum (2024) investigated", "Orth and Baum [37] investigated"),
    ("Rakhewar et al. (2023): In their work, they", "Rakhewar et al. [38]: In their work, they"),
    ("Yadav et al. (2023) ,their work focuses", "Yadav et al. [39], whose work focuses"),
    ("Sehgal & Yathrath (2022): In their work, they", "Sehgal and Yathrath [40]: In their work, they"),
    ("Chatterjee et al. (2021) present a comprehensive", "Chatterjee et al. [41] present a comprehensive"),
    ("Vallas and Schor (2020) provide a conceptual", "Vallas and Schor [42] provide a conceptual"),
    ("Rana et al. (2019): In their work, they contributed", "Rana et al. [7]: In their work, they contributed"),
    ("Indravasan et al. (2018) made a significant", "Indravasan et al. [8] made a significant"),
    ("Berg, E. Rani (2018),", "Berg [43],"),
    ("Sundararajan (2016) examined the rise", "Sundararajan [44] examined the rise"),

    # Ch3 - Design
    ("A Use Case Diagram illustrates the primary interactions between the system actors and the system itself.",
     "A Use Case Diagram [50] illustrates the primary interactions between the system actors and the system itself."),
    ("A Context Diagram (also known as a Level 0 Data Flow Diagram) is the highest-level abstraction of a system.",
     "A Context Diagram (also known as a Level 0 Data Flow Diagram) [51] is the highest-level abstraction of a system."),
    ("A Data Flow Diagram (DFD) is a graphical and logical representation that illustrates how data moves through a system,",
     "A Data Flow Diagram (DFD) is a graphical and logical representation that illustrates how data moves through a system [52],"),
    ("An Entity Relationship Diagram (ERD) is a conceptual blueprint of your database.",
     "An Entity Relationship Diagram (ERD) is a conceptual blueprint of the database [25]."),
    ("The Relational Data Model (RDM), also known as the Physical Database Schema, is the most technical version of your data design.",
     "The Relational Data Model (RDM), also known as the Physical Database Schema [53], is the most technical version of the data design."),
]

total_refs = 0
for old, new in refs:
    c = force_replace(doc, old, new)
    if c > 0:
        total_refs += c

print(f"  Total reference insertions: {total_refs}")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_FILE)
print(f"\n✅ Document saved to: {OUTPUT_FILE}")

# ============================================================
# VERIFICATION
# ============================================================
import subprocess, re
result = subprocess.run(["pandoc", OUTPUT_FILE, "-t", "plain"], capture_output=True, text=True)
text = result.stdout

print("\n=== VERIFICATION ===")
tech_terms = ["Next.js 16", "TypeScript", "Prisma ORM", "SQLite", "Bun", "JWT", "Paystack", "Vercel", "Socket.io"]
for t in tech_terms:
    found = t in text
    print(f"  {'✅' if found else '❌'} {t}")

old_terms = ["Node.js with Express", "Firebase is used", "HTML5, CSS3, and JavaScript,for", "MongoDB(NOSQL)", "Stripe or Paystack"]
for t in old_terms:
    found = t in text
    print(f"  {'✅ Removed' if not found else '⚠️ Still present'}: {t}")

# Count citations
all_refs = re.findall(r'\[\d+\]', text)
print(f"\n  Total inline citations: {len(all_refs)}")
unique = sorted(set(int(r.strip('[]')) for r in all_refs))
print(f"  Unique citation numbers: {unique}")

# Check page count estimate  
word_count = len(text.split())
print(f"  Approx word count: {word_count}")
