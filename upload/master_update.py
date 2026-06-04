"""
MASTER SCRIPT: Apply ALL changes to HomeEase writeup in one pass.
Reads original, applies tech stack + inline references, saves updated.
"""
from docx import Document
import subprocess, re

INPUT = "/home/z/my-project/upload/new mariam.docx"
OUTPUT = "/home/z/my-project/upload/HomeEase_Project_Writeup.docx"

doc = Document(INPUT)

def replace_para_text(para, old, new):
    """Replace text in paragraph, handling multi-run splits."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_full)
    return True

def apply_all(doc, replacements):
    """Apply replacements to all paragraphs and table cells."""
    count = 0
    for para in doc.paragraphs:
        for old, new in replacements:
            if replace_para_text(para, old, new):
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if replace_para_text(para, old, new):
                            count += 1
    return count

# ============================================================
# ALL REPLACEMENTS IN ORDER (tech stack + references combined)
# ============================================================
replacements = [
    # ---- SECTION 1.E: THE BIG METHODOLOGY PARAGRAPH ----
    (
        "The system implementation follows an agile development methodology, enabling incremental development and continuous feedback. The front end of the virtual space is developed using HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces.  The backend system is implemented using Node.js with Express which manage business logic, service scheduling, user authentication, and communication between system components. Data is stored and managed using a relational database such as MySQL for scalability and data structure requirements.",
        "The system implementation follows an agile development methodology [31], enabling incremental development and continuous feedback. The front end of the virtual space is developed using Next.js 16 with TypeScript and Tailwind CSS 4 for type-safe, responsive, and dynamic user interfaces. The backend system is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, which manages business logic, service scheduling, user authentication, real-time messaging via Socket.io, and communication between system components. Data is stored and managed using SQLite as the relational database engine, with Prisma ORM providing type-safe schema management and an intuitive query API for data access."
    ),
    (
        "Firebase is used for secure user authentication and role-based access control, while Stripe or Paystack APIs are employed for secure payment processing.",
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication and role-based access control, while the Paystack payment gateway API is employed for secure payment processing integrated with a virtual wallet and escrow system."
    ),
    (
        "Cloud services such as Google Cloud, or Hostinger support hosting, storage, and scalability of the virtual space.",
        "Vercel, a cloud platform optimized for Next.js applications, supports hosting, serverless function execution, and automatic scalability of the virtual space."
    ),
    (
        "software testing tools such as Jest, Selenium, or Postman to verify functionality, performance, and security.",
        "software testing tools such as Jest for unit testing and Postman for API endpoint verification to verify functionality, performance, and security."
    ),

    # ---- SECTION 3.4 INTRO ----
    (
        "Commonly used DBMS solutions such as MySQL or MongoDB are suitable due to their reliability and scalability.",
        "SQLite is used as the relational database engine, managed through Prisma ORM which provides type-safe database access and schema management, ensuring reliability and data integrity."
    ),

    # ---- SECTION 3.4.1 TABLE (cell by cell) ----
    ("MongoDB(NOSQL)", "SQLite with Prisma ORM"),
    ("React native +", "Next.js 16 + TypeScript +"),
    ("Node.js with", "Next.js 16 (App Router) +"),
    ("Paystack or", "Paystack Payment"),
    ("Flutterwave", "Gateway"),

    # ---- CHAPTER 1: INLINE REFERENCES ----
    ("The rapid advancement of digital technologies has transformed how services are delivered across various industries. Domestic",
     "The rapid advancement of digital technologies has transformed how services are delivered across various industries [1]. Domestic"),
    ("Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping. Although",
     "Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping [2]. Although"),
    ("The concept of a virtual space for domestic service involves creating an online platform or digital environment",
     "The concept of a virtual space for domestic service involves creating an online platform or digital environment [3]"),
    ("Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance, developing",
     "Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance [4], developing"),
    ("The provision of domestic services is largely managed through informal and manual processes, resulting",
     "The provision of domestic services is largely managed through informal and manual processes [5, 6], resulting"),
    ("Khatri & Gupta (2020)", "Khatri and Gupta [5]"),
    ("Adeyemi & Fatile\n                (2021)", "Adeyemi and Fatile\n                [6]"),
    ("Rana et al.\n   (2019)", "Rana et al.\n   [7]"),
    ("Indravasan\n   et al.\n                (2018)", "Indravasan\n   et al.\n                [8]"),
    ("Chen et al. (2021)", "Chen et al. [9]"),

    # ---- CHAPTER 2: THEORETICAL FRAMEWORK ----
    ("perceived ease of use (Davis, 1989). In the context",
     "perceived ease of use [10]. In the context"),
    ("The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries",
     "The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries [11]"),
    ("Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms. Dimensions",
     "Additionally, Service Quality Theory [12] explains how users evaluate services delivered through digital platforms. Dimensions"),

    # ---- CHAPTER 2: VIRTUAL SPACE / PAYMENTS ----
    ("When integrated with Paystack digital payment system, the",
     "When integrated with a digital payment system such as Paystack [14], the"),
    ("Digital payment systems have achieved widespread adoption due to their",
     "Digital payment systems have achieved widespread adoption due to their [14]"),

    # ---- CHAPTER 2: WEB 3.0 ----
    ("Web 3.0 refers to the next stage in the evolution of the internet, designed to move",
     "Web 3.0 refers to the next stage in the evolution of the internet [15], designed to move"),
    ("The heart of Web 3.0 is blockchain technology. A blockchain is a distributed ledger",
     "The heart of Web 3.0 is blockchain technology. A blockchain is a distributed ledger [16]"),

    # ---- CHAPTER 2: DATABASE ----
    ("A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated. Databases are",
     "A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated [17]. Databases are"),
    ("ACID properties are fundamental principles that ensure the reliability and correctness of database transactions. A transaction refers",
     "ACID properties are fundamental principles that ensure the reliability and correctness of database transactions [20]. A transaction refers"),

    # ---- CHAPTER 2: DATA MODELING ----
    ("Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system. It defines",
     "Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system [24]. It defines"),
    ("Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures. It is widely",
     "Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures [25]. It is widely"),
    ("Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency. It involves",
     "Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency [26]. It involves"),

    # ---- CHAPTER 2: SDLC ----
    ("A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems. It provides",
     "A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems [27]. It provides"),
    ("System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner. It provides a step-by-step approach",
     "System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner [28]. It provides a step-by-step approach"),
    ("The waterfall model is a sequential, plan driven-process where you must",
     "The waterfall model is a sequential, plan-driven process [29] where you must"),
    ("The V model (Verification and Validation model) is an extension of the waterfall model. All the requirements are gathered at the start",
     "The V model (Verification and Validation model) is an extension of the waterfall model [30]. All the requirements are gathered at the start"),
    ("The agile process model encourages continuous iterations of development and testing.",
     "The agile process model [31] encourages continuous iterations of development and testing."),

    # ---- CHAPTER 2: HOSTING ----
    ("Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure. These infrastructures",
     "Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure [32]. These infrastructures"),
    ("Platform as a Service (PaaS) provides a complete hosting environment that includes operating systems, development frameworks, databases, and deployment tools.",
     "Platform as a Service (PaaS) [33] provides a complete hosting environment that includes operating systems, development frameworks, databases, and deployment tools."),

    # ---- CHAPTER 2: RELATED WORKS ----
    ("Aishwaryalakshmi et al. (2024), focuses on", "Aishwaryalakshmi et al. [34], focuses on"),
    ("Pais and Zanoni (2024), Their study adopted", "Pais and Zanoni [35]. Their study adopted"),
    ("Meyanban et al. (2024), This study explored", "Meyanban et al. [36]. This study explored"),
    ("Orth and Baum (2024) investigated", "Orth and Baum [37] investigated"),
    ("Rakhewar et al. (2023): In their work, they", "Rakhewar et al. [38]: In their work, they"),
    ("Yadav et al. (2023) ,their work focuses", "Yadav et al. [39], whose work focuses"),
    ("Sehgal & Yathrath (2022): In their work, they", "Sehgal and Yathrath [40]: In their work, they"),
    ("Chatterjee et al. (2021) present a comprehensive", "Chatterjee et al. [41] present a comprehensive"),
    ("Vallas and Schor (2020) provide a conceptual", "Vallas and Schor [42] provide a conceptual"),
    ("Rana et al. (2019): In their work, they contributed", "Rana et al. [7]: In their work, they contributed"),
    ("Indravasan et al. (2018) made a significant", "Indravasan et al. [8] made a significant"),
    ("Berg, E. Rani (2018),", "Berg [43],"),
    ("Sundararajan (2016) examined the rise", "Sundararajan [44] examined the rise"),

    # ---- CHAPTER 3: DIAGRAMS ----
    ("A Use Case Diagram illustrates the primary interactions between the system actors and the system itself.",
     "A Use Case Diagram [50] illustrates the primary interactions between the system actors and the system itself."),
    ("A Context Diagram (also known as a Level 0 Data Flow Diagram) is the highest-level abstraction of a system.",
     "A Context Diagram (also known as a Level 0 Data Flow Diagram) [51] is the highest-level abstraction of a system."),
    ("A Data Flow Diagram (DFD) is a graphical and logical representation that illustrates how data moves through a system,",
     "A Data Flow Diagram (DFD) is a graphical and logical representation that illustrates how data moves through a system [52],"),
    ("An Entity Relationship Diagram (ERD) is a conceptual blueprint of your database.",
     "An Entity Relationship Diagram (ERD) is a conceptual blueprint of the database [25]."),
    ("The Relational Data Model (RDM), also known as the Physical Database Schema, is the most technical version of your data design.",
     "The Relational Data Model (RDM), also known as the Physical Database Schema [53], is the most technical version of the data design."),
]

print("=== Applying ALL Replacements (Tech Stack + References) ===")
total = apply_all(doc, replacements)
print(f"Total replacements applied: {total}")

doc.save(OUTPUT)
print(f"\nSaved to: {OUTPUT}")

# ============================================================
# COMPREHENSIVE VERIFICATION  
# ============================================================
result = subprocess.run(["pandoc", OUTPUT, "-t", "plain"], capture_output=True, text=True)
text = result.stdout

print("\n=== TECH STACK VERIFICATION ===")
for t in ["Next.js 16", "TypeScript", "Tailwind CSS 4", "Bun", "Prisma ORM", "SQLite", "Socket.io", "JWT", "Paystack", "Vercel"]:
    print(f"  {'✅' if t in text else '❌'} {t}")

print("\n=== OLD STACK REMOVAL CHECK ===")
for t in ["HTML5, CSS3, and JavaScript,for", "Node.js with Express which manage", "Firebase is used", "Stripe or Paystack APIs", "MongoDB(NOSQL)"]:
    print(f"  {'✅ Removed' if t not in text else '⚠️ Still present'}: {t[:50]}")

print("\n=== CITATIONS ===")
all_citations = re.findall(r'\[\d+\]', text)
unique_nums = sorted(set(int(c.strip('[]')) for c in all_citations))
print(f"  Total citations: {len(all_citations)}")
print(f"  Unique ref numbers: {unique_nums}")

wc = len(text.split())
print(f"\n  Word count: ~{wc}")
print(f"  Estimated pages: ~{wc // 300} (at ~300 words/page)")
