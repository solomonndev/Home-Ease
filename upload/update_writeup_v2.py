"""
Robust script to update HomeEase project writeup:
1. Update tech stack to match actual implementation (handles multi-run text)
2. Add inline academic references
"""
from docx import Document
from copy import deepcopy
import re

INPUT_FILE = "/home/z/my-project/upload/new mariam.docx"
OUTPUT_FILE = "/home/z/my-project/upload/HomeEase_Project_Writeup.docx"

doc = Document(INPUT_FILE)

def robust_replace_in_para(para, old_text, new_text):
    """Replace text in a paragraph even if split across runs."""
    full_text = "".join(r.text for r in para.runs)
    if old_text not in full_text:
        return False
    
    new_full = full_text.replace(old_text, new_text)
    if new_full == full_text:
        return False
    
    # Strategy: find which runs contain the old text and rebuild
    # Simple approach: put all text in first run, clear rest
    # But we lose formatting. Better: preserve run formatting proportionally.
    
    # Find the start position of old_text in concatenated text
    start_idx = full_text.find(old_text)
    end_idx = start_idx + len(old_text)
    new_len = len(new_text)
    
    # Build new text distribution across runs
    # Walk through runs, tracking character positions
    new_runs_text = []
    pos = 0
    replacement_done = False
    
    for run in para.runs:
        run_len = len(run.text)
        run_start = pos
        run_end = pos + run_len
        
        if not replacement_done:
            # Check if this run overlaps with the replacement zone
            if run_end <= start_idx:
                # Run is entirely before replacement - keep as is
                new_runs_text.append(run.text)
            elif run_start >= end_idx:
                # Run is entirely after replacement
                new_runs_text.append(run.text)
                replacement_done = True
            else:
                # Run overlaps with replacement zone
                before = run.text[:max(0, start_idx - run_start)]
                after = run.text[min(run_len, end_idx - run_start):]
                
                if run_start < start_idx and run_end > end_idx:
                    # Run fully contains the replacement
                    new_runs_text.append(before + new_text + after)
                    replacement_done = True
                elif run_start < start_idx:
                    # Run has the beginning of old text
                    new_runs_text.append(before + new_text)
                    if run_end >= end_idx:
                        replacement_done = True
                elif run_end <= end_idx:
                    # Run is within old text - will be handled or already empty
                    if run_end >= end_idx:
                        new_runs_text.append(after)
                        replacement_done = True
                    else:
                        new_runs_text.append("")
                else:
                    # Run has the end of old text
                    new_runs_text.append(after)
                    replacement_done = True
        else:
            new_runs_text.append(run.text)
        
        pos = run_end
    
    # Apply new texts to runs
    for i, run in enumerate(para.runs):
        if i < len(new_runs_text):
            run.text = new_runs_text[i]
    
    return True


def replace_in_all_paragraphs(doc, replacements):
    """Apply replacements to all paragraphs and table cells."""
    count = 0
    # Process paragraphs
    for para in doc.paragraphs:
        for old_text, new_text in replacements:
            if robust_replace_in_para(para, old_text, new_text):
                count += 1
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements:
                        if robust_replace_in_para(para, old_text, new_text):
                            count += 1
    
    return count

# ============================================================
# TECH STACK REPLACEMENTS (ordered from longest to shortest for safety)
# ============================================================

tech_replacements = [
    # Section 1.E Methodology - main paragraph
    (
        "The front end of the virtual space is developed using HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces. The backend system is implemented using Node.js with Express which manage business logic, service scheduling, user authentication, and communication between system components. Data is stored and managed using a relational database such as MySQL for scalability and data structure requirements.",
        "The front end of the virtual space is developed using Next.js 16 with TypeScript and Tailwind CSS 4 for type-safe, responsive, and dynamic user interfaces. The backend system is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, which manages business logic, service scheduling, user authentication, real-time messaging via Socket.io, and communication between system components. Data is stored and managed using SQLite as the relational database engine, with Prisma ORM providing type-safe schema management and an intuitive query API for data access."
    ),
    (
        "Firebase is used for secure user authentication and role-based access control, while Stripe or Paystack APIs are employed for secure payment processing.",
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication and role-based access control, while the Paystack payment gateway API is employed for secure payment processing integrated with a virtual wallet and escrow system."
    ),
    (
        "Cloud services such as Google Cloud, or Hostinger support hosting, storage, and scalability of the virtual space.",
        "Vercel, a cloud platform optimized for Next.js applications, supports hosting, serverless function execution, and automatic scalability of the virtual space."
    ),
    (
        "software testing tools such as Jest, Selenium, or Postman",
        "software testing tools such as Jest for unit testing and Postman for API endpoint verification"
    ),
    (
        "The system implementation follows an agile development methodology, enabling incremental development and continuous feedback.",
        "The system implementation follows an agile development methodology [31], enabling incremental development and continuous feedback."
    ),

    # Section 3.4 intro - DBMS reference
    (
        "Commonly used DBMS solutions such as MySQL or MongoDB are suitable due to their reliability and scalability.",
        "SQLite is used as the relational database engine, managed through Prisma ORM which provides type-safe database access and schema management, ensuring reliability and data integrity."
    ),
]

print("Applying tech stack replacements...")
tech_count = replace_in_all_paragraphs(doc, tech_replacements)
print(f"  Tech stack replacements applied: {tech_count}")

# Also fix table cells for Software Requirements table
# The table replacements need cell-level approach
table_fixes = {
    "MongoDB(NOSQL)": "SQLite with Prisma ORM",
    "MongoDB (NOSQL)": "SQLite with Prisma ORM", 
    "React native +": "Next.js 16 + TypeScript +",
    "Node.js with": "Next.js 16 (App Router) +",
    "Paystack or": "Paystack Payment",
    "Flutterwave": "Gateway",
}

table_count = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full_text = "".join(r.text for r in para.runs)
                modified = False
                for old_t, new_t in table_fixes.items():
                    if old_t in full_text:
                        full_text = full_text.replace(old_t, new_t)
                        modified = True
                if modified:
                    for r in para.runs:
                        r.text = ""
                    if para.runs:
                        para.runs[0].text = full_text
                    else:
                        para.add_run(full_text)
                    table_count += 1

print(f"  Table cell fixes: {table_count}")

# ============================================================
# INLINE REFERENCES
# ============================================================

inline_refs = [
    # Chapter 1 - Background
    ("The rapid advancement of digital technologies has transformed how services are delivered across various industries.", 
     "The rapid advancement of digital technologies has transformed how services are delivered across various industries [1]."),
    ("Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping.",
     "Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping [2]."),
    ("The concept of a virtual space for domestic service involves creating an online platform",
     "The concept of a virtual space for domestic service involves creating an online platform [3]"),
    ("Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance",
     "Given the growing reliance on digital platforms and the increasing need for organized access to domestic assistance [4]"),

    # Chapter 1 - Problem Statement
    ("The provision of domestic services is largely managed through informal and manual processes",
     "The provision of domestic services is largely managed through informal and manual processes [5, 6]"),

    # Chapter 1 - Lit Review table
    ("Khatri &", "Khatri and"),
    ("(Khatri & Gupta (2020)", "(Khatri and Gupta [5]"),
    ("(Adeyemi & Fatile", "(Adeyemi and Fatile [6]"),
    ("(Rana et al. (2019)", "(Rana et al. [7]"),
    ("(Indravasan et al. (2018)", "(Indravasan et al. [8]"),
    ("(Chen et al. (2021)", "(Chen et al. [9]"),

    # Chapter 2 - TAM
    ("(Davis, 1989).", "[10]."),
    
    # Chapter 2 - Platform Economy
    ("The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries",
     "The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries [11]"),
    ("Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms.",
     "Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms [12]."),
    ("Service Quality Theory",
     "Service Quality Theory [12]"),

    # Chapter 2 - Virtual Space section
    ("When integrated with Paystack digital payment system",
     "When integrated with a digital payment system such as Paystack [14]"),
    ("Digital payment systems have achieved widespread adoption",
     "Digital payment systems have achieved widespread adoption [14]"),

    # Chapter 2 - Web 3.0
    ("Web 3.0 refers to the next stage in the evolution of the internet",
     "Web 3.0 refers to the next stage in the evolution of the internet [15]"),
    ("blockchain is a distributed ledger",
     "blockchain is a distributed ledger [16]"),

    # Chapter 2 - Database  
    ("A database is an organized collection of data that is stored electronically",
     "A database is an organized collection of data that is stored electronically [17]"),
    ("ACID properties are fundamental principles",
     "ACID properties are fundamental principles [20]"),

    # Chapter 2 - Data Modeling
    ("Data modeling refers to the process of creating a visual and logical representation",
     "Data modeling refers to the process of creating a visual and logical representation [24]"),
    ("Entity\u2013Relationship (ER) Modeling",
     "Entity\u2013Relationship (ER) Modeling [25]"),
    ("Normalization: Normalization is the process",
     "Normalization: Normalization is the process [26]"),

    # Chapter 2 - SDLC
    ("A software process model is a structured framework",
     "A software process model is a structured framework [27]"),
    ("System Development Life Cycle (SDLC) is a systematic framework",
     "System Development Life Cycle (SDLC) is a systematic framework [28]"),
    ("The waterfall model is a sequential, plan driven-process",
     "The waterfall model is a sequential, plan-driven process [29]"),
    ("The V model (Verification and Validation model)",
     "The V model (Verification and Validation model) [30]"),
    ("The agile process model encourages continuous iterations",
     "The agile process model encourages continuous iterations [31]"),

    # Chapter 2 - Hosting
    ("Software hosting is a fundamental concept in modern software engineering",
     "Software hosting is a fundamental concept in modern software engineering [32]"),
    ("Platform as a Service (PaaS) provides a complete hosting environment",
     "Platform as a Service (PaaS) provides a complete hosting environment [33]"),

    # Chapter 2 - Related Works
    ("Aishwaryalakshmi et al. (2024),", "Aishwaryalakshmi et al. [34],"),
    ("Pais and Zanoni (2024),", "Pais and Zanoni [35],"),
    ("Meyanban et al. (2024),", "Meyanban et al. [36],"),
    ("Orth and Baum (2024)", "Orth and Baum [37]"),
    ("Rakhewar et al. (2023):", "Rakhewar et al. [38]:"),
    ("Yadav et al. (2023)", "Yadav et al. [39]"),
    ("Sehgal & Yathrath (2022):", "Sehgal and Yathrath [40]:"),
    ("Chatterjee et al. (2021)", "Chatterjee et al. [41]"),
    ("Vallas and Schor (2020)", "Vallas and Schor [42]"),
    ("Sundararajan (2016)", "Sundararajan [44]"),

    # Chapter 3 - Use Case
    ("A Use Case Diagram illustrates the primary interactions",
     "A Use Case Diagram illustrates the primary interactions [50]"),
    ("A Context Diagram (also known as a Level 0 Data Flow Diagram)",
     "A Context Diagram (also known as a Level 0 Data Flow Diagram) [51]"),
    ("A Data Flow Diagram (DFD) is a graphical and logical representation",
     "A Data Flow Diagram (DFD) is a graphical and logical representation [52]"),
    ("An Entity Relationship Diagram (ERD) is a conceptual blueprint",
     "An Entity Relationship Diagram (ERD) is a conceptual blueprint [25]"),
    ("The Relational Data Model (RDM), also known as the Physical Database Schema",
     "The Relational Data Model (RDM), also known as the Physical Database Schema [53]"),
]

print("\nApplying inline references...")
ref_count = replace_in_all_paragraphs(doc, inline_refs)
print(f"  Inline references applied: {ref_count}")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_FILE)
print(f"\nDocument saved to: {OUTPUT_FILE}")

# Verify
import subprocess
result = subprocess.run(
    ["pandoc", OUTPUT_FILE, "-t", "plain"],
    capture_output=True, text=True
)
text = result.stdout
lines = text.split("\n")

# Check tech stack
print("\n--- Verification ---")
for term in ["Next.js 16", "TypeScript", "Tailwind CSS 4", "Prisma ORM", "SQLite", "Bun", "JWT", "Paystack", "Vercel", "shadcn/ui"]:
    found = term in text
    print(f"  {'✅' if found else '❌'} {term}: {'Found' if found else 'NOT FOUND'}")

for term in ["Node.js with Express", "Firebase", "HTML5, CSS3", "MongoDB(NOSQL)"]:
    found = term in text
    status = "Still present (check context)" if found else "Removed"
    print(f"  {'⚠️' if found else '✅'} Old: '{term}': {status}")

# Count references
import re
refs = re.findall(r'\[\d+\]', text)
unique_refs = set(refs)
print(f"\n  Total inline citations found: {len(refs)}")
print(f"  Unique reference numbers: {sorted([int(r.strip('[]')) for r in unique_refs])}")
