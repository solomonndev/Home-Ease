#!/usr/bin/env python3
"""
Generate Chapter 4 — System Implementation and User Guide
HomeEase project with 22 screenshots.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, "screenshots")

FIGURES = {
    1:  os.path.join(IMG_DIR, "fig01_hero.png"),
    2:  os.path.join(IMG_DIR, "fig03_services.png"),
    3:  os.path.join(IMG_DIR, "fig04_features.png"),
    4:  os.path.join(IMG_DIR, "fig05_howitworks.png"),
    5:  os.path.join(IMG_DIR, "fig07_roleselect.png"),
    6:  os.path.join(IMG_DIR, "fig08_seeker_signup.png"),
    7:  os.path.join(IMG_DIR, "fig10_provider_top.png"),
    8:  os.path.join(IMG_DIR, "fig11_provider_skills.png"),
    9:  os.path.join(IMG_DIR, "fig12_provider_bank.png"),
    10: os.path.join(IMG_DIR, "fig09_login.png"),
    11: os.path.join(BASE, "Screenshot (20).png"),
    12: os.path.join(BASE, "Screenshot (21).png"),
    13: os.path.join(BASE, "Screenshot (27).png"),
    14: os.path.join(BASE, "Screenshot (19).png"),
    15: os.path.join(BASE, "Screenshot (28).png"),
    16: os.path.join(BASE, "Screenshot (24).png"),
    17: os.path.join(BASE, "Screenshot (39).png"),
    18: os.path.join(BASE, "Screenshot (40).png"),
    19: os.path.join(BASE, "Screenshot (33).png"),
    20: os.path.join(BASE, "Screenshot (34).png"),
    21: os.path.join(BASE, "Screenshot (35).png"),
    22: os.path.join(IMG_DIR, "fig13_arch.png"),
}

CAPTIONS = {
    1:  "HomeEase Landing Page Hero Section",
    2:  "HomeEase Services Grid Display",
    3:  "HomeEase Platform Features Overview",
    4:  "HomeEase How It Works Section",
    5:  "Role Selection – Service Seeker or Service Provider",
    6:  "Service Seeker Registration Form",
    7:  "Service Provider Registration – Personal Information",
    8:  "Service Provider Registration – Skills and Services",
    9:  "Service Provider Registration – Bank Details",
    10: "HomeEase Login Form",
    11: "Find Artisans – Search Page",
    12: "Find Artisans – Search Results",
    13: "My Requests List",
    14: "Payment Page",
    15: "My Requests with Review and Rating",
    16: "Messages / Chat Screen",
    17: "Provider Verification Pending Screen",
    18: "Provider Chat with Admin During Verification",
    19: "Provider My Jobs – Checking In and Out",
    20: "Admin Payouts Screen",
    21: "Admin Payouts Confirmation",
    22: "HomeEase System Architecture Diagram",
}

# ──────────────────────────────────────────────
# Helper functions
# ──────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set background colour of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """Apply borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows):
    """Create a formatted table with light-blue headers and borders."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        set_cell_shading(cell, "D6EAF8")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
    return table


def add_body(doc, text):
    """Add a normal body paragraph (12pt, Times New Roman, 1.5 spacing)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_body_rich(doc, parts):
    """Add a paragraph with mixed formatting. parts is a list of (text, bold) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.bold = bold
    return p


def add_step(doc, step_num, text):
    """Add a numbered step with bold 'Step N:' prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run_label = p.add_run(f"Step {step_num}: ")
    run_label.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = "Times New Roman"
    run_text = p.add_run(text)
    run_text.font.size = Pt(12)
    run_text.font.name = "Times New Roman"
    return p


def add_bullet(doc, text):
    """Add a bullet point using the bullet character."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(f"\u2022  {text}")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_heading_level1(doc, text):
    """Section heading — 14pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    return p


def add_heading_level2(doc, text):
    """Sub-section heading — 12pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_figure(doc, fig_num):
    """Insert a figure image (centered, 5.5 in wide) and caption below it."""
    path = FIGURES[fig_num]
    caption_text = f"Figure 4.{fig_num}: {CAPTIONS[fig_num]}"
    # Image
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)
    run_img = p_img.add_run()
    run_img.add_picture(path, width=Inches(5.5))
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(caption_text)
    run_cap.italic = True
    run_cap.font.size = Pt(10)
    run_cap.font.name = "Times New Roman"


def add_page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────

def main():
    doc = Document()

    # ── Page setup: 1-inch margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Default style ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.5

    # ═══════════════════════════════════════════
    # CHAPTER TITLE
    # ═══════════════════════════════════════════
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    run_ch = p_title.add_run("CHAPTER 4")
    run_ch.bold = True
    run_ch.font.size = Pt(16)
    run_ch.font.name = "Times New Roman"

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("SYSTEM IMPLEMENTATION AND USER GUIDE")
    run_sub.bold = True
    run_sub.font.size = Pt(16)
    run_sub.font.name = "Times New Roman"

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.1 Introduction
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.1 Introduction")
    add_body(doc,
        "This chapter provides a comprehensive walkthrough of the HomeEase platform, "
        "detailing the system implementation environment and serving as a complete user "
        "manual for all three user roles: Service Seekers, Service Providers, and "
        "Administrators. Each section is accompanied by annotated screenshots that illustrate "
        "the key interfaces and workflows of the system. The chapter covers the hardware and "
        "software requirements, step-by-step instructions for every major feature, the "
        "underlying system architecture, and a sample end-to-end system run demonstrating a "
        "typical service booking scenario."
    )
    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.2 Implementation Environment
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.2 Implementation Environment")

    # 4.2.1 Hardware Requirements
    add_heading_level2(doc, "4.2.1 Hardware Requirements")
    add_body(doc, "Table 4.1 presents the hardware specifications used during the development of the HomeEase platform.")
    doc.add_paragraph()  # spacer
    p_tbl_cap1 = doc.add_paragraph()
    p_tbl_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tbl_cap1.add_run("Table 4.1: Development Hardware Requirements")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"

    make_table(doc,
        ["Component", "Specification"],
        [
            ["Processor", "Intel Core i5 / AMD Ryzen 5 or higher"],
            ["RAM", "8 GB minimum (16 GB recommended)"],
            ["Hard Disk", "256 GB SSD minimum"],
            ["Display", "1366 \u00d7 768 resolution minimum"],
            ["Network", "Broadband internet connection"],
        ]
    )
    add_body(doc, "")
    add_body(doc, "Table 4.2 shows the recommended hardware for the production server deployment.")
    p_tbl_cap2 = doc.add_paragraph()
    p_tbl_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tbl_cap2.add_run("Table 4.2: Server Hardware Requirements")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"

    make_table(doc,
        ["Component", "Specification"],
        [
            ["Processor", "Intel Xeon / AMD EPYC (multi-core)"],
            ["RAM", "16 GB minimum (32 GB recommended)"],
            ["Hard Disk", "100 GB SSD + scalable storage"],
            ["Network", "Dedicated 100 Mbps+ bandwidth"],
            ["Operating System", "Ubuntu Server 22.04 LTS"],
        ]
    )
    add_body(doc, "")

    # 4.2.2 Software Requirements
    add_heading_level2(doc, "4.2.2 Software Requirements")
    add_body(doc, "Table 4.3 lists the software tools and technologies used throughout the implementation of the HomeEase platform.")
    doc.add_paragraph()
    p_tbl_cap3 = doc.add_paragraph()
    p_tbl_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tbl_cap3.add_run("Table 4.3: Software Requirements")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"

    make_table(doc,
        ["Software / Technology", "Version / Details", "Purpose"],
        [
            ["React.js", "v18.x", "Frontend framework"],
            ["Node.js", "v18.x LTS", "Backend runtime environment"],
            ["Express.js", "v4.x", "Backend web framework"],
            ["MongoDB", "v6.x", "NoSQL database"],
            ["Mongoose ODM", "v7.x", "Object data modelling for MongoDB"],
            ["Tailwind CSS", "v3.x", "Utility-first CSS framework"],
            ["Socket.IO", "v4.x", "Real-time bidirectional communication"],
            ["Flutterwave", "v3.x API", "Payment processing gateway"],
            ["Cloudinary", "SDK", "Cloud image storage and management"],
            ["VS Code", "Latest", "Integrated development environment"],
            ["Postman", "Latest", "API testing and documentation"],
            ["Git / GitHub", "Latest", "Version control"],
            ["npm / Yarn", "v9.x / v1.22.x", "Package management"],
        ]
    )
    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.3 System Interface and User Manual
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.3 System Interface and User Manual")
    add_body(doc,
        "This section provides a detailed walkthrough of every interface and feature "
        "available on the HomeEase platform. Instructions are organised by user role and "
        "cover the complete workflow from registration to service completion."
    )

    # ── 4.3.1 Getting Started ──
    add_heading_level2(doc, "4.3.1 Getting Started (Common to All Users)")

    # 4.3.1.1 Accessing the Website
    add_heading_level2(doc, "4.3.1.1 Accessing the Website")
    add_step(doc, 1, 'Open any modern web browser such as Google Chrome, Mozilla Firefox, or Microsoft Edge.')
    add_step(doc, 2, "Enter the HomeEase platform URL in the browser's address bar and press **Enter**.")
    add_step(doc, 3, "The landing page loads, displaying the hero section with a brief overview of the platform, as shown in Figure 4.1.")
    add_figure(doc, 1)

    # 4.3.1.2 Platform Services and Features
    add_heading_level2(doc, "4.3.1.2 Platform Services and Features")
    add_body(doc,
        "After the hero section, users can scroll down to view the full range of services "
        "offered by the platform. The **Services** section presents a visual grid of service "
        "categories, allowing users to quickly identify the type of artisan they need. "
        "Each category is represented by an icon and label, as illustrated in Figure 4.2."
    )
    add_figure(doc, 2)

    add_body(doc,
        "Below the services grid, the **Features** section highlights the key benefits of "
        "using the HomeEase platform. Features such as verified artisans, secure payments, "
        "real-time messaging, and ratings and reviews are presented in an easy-to-scan card "
        "layout, as shown in Figure 4.3."
    )
    add_figure(doc, 3)

    add_body(doc,
        "The **How It Works** section provides a simplified three-step guide for new users. "
        "It explains the basic workflow of searching for a service provider, booking a "
        "service, and completing the transaction. This section gives first-time visitors a "
        "clear understanding of the platform's core process, as depicted in Figure 4.4."
    )
    add_figure(doc, 4)

    # 4.3.1.3 Creating an Account
    add_heading_level2(doc, "4.3.1.3 Creating an Account")
    add_body(doc,
        "HomeEase supports two primary user roles: **Service Seeker** and **Service Provider**. "
        "The registration process begins with selecting the appropriate role."
    )

    add_step(doc, 1, "On the landing page, click the **Get Started** or **Sign Up** button to initiate the registration process.")
    add_step(doc, 2, "The system presents a role selection screen, as shown in Figure 4.5. Choose **Service Seeker** if you need to hire artisans, or **Service Provider** if you offer services.")
    add_figure(doc, 5)

    add_body(doc, "Service Seeker Registration:")
    add_step(doc, 3, "After selecting **Service Seeker**, the registration form appears. Fill in the required fields: **Full Name**, **Email Address**, **Phone Number**, and **Password**.")
    add_step(doc, 4, "Click the **Sign Up** or **Create Account** button to complete registration, as shown in Figure 4.6.")
    add_figure(doc, 6)

    add_body(doc, "Service Provider Registration:")
    add_step(doc, 5, "After selecting **Service Provider**, a multi-section registration form is displayed. Begin by entering personal details: **Full Name**, **Email Address**, **Phone Number**, **Password**, and **Location**, as shown in Figure 4.7.")
    add_figure(doc, 7)
    add_step(doc, 6, "Scroll down to the **Skills and Services** section. Select the service categories you offer from the available options (e.g., Plumbing, Electrical, Cleaning, Painting). You may also add a brief description of your experience, as illustrated in Figure 4.8.")
    add_figure(doc, 8)
    add_step(doc, 7, "Proceed to the **Bank Details** section. Enter your **Bank Name**, **Account Number**, and **Account Name** to enable payout processing once you start earning, as shown in Figure 4.9.")
    add_figure(doc, 9)
    add_step(doc, 8, "Review all entered information and click **Submit** to complete the registration. New providers are placed in a **pending verification** state until an administrator reviews and approves their profile.")

    # 4.3.1.4 Signing In
    add_heading_level2(doc, "4.3.1.4 Signing In")
    add_step(doc, 1, "On the landing page, click the **Login** button located in the navigation bar.")
    add_step(doc, 2, "Enter your registered **Email Address** and **Password** in the login form, as shown in Figure 4.10.")
    add_figure(doc, 10)
    add_step(doc, 3, 'Click the **Login** button. Upon successful authentication, the system redirects you to the appropriate dashboard based on your user role (Seeker, Provider, or Admin).')

    # ── 4.3.2 Service Seeker User Manual ──
    add_heading_level2(doc, "4.3.2 Service Seeker User Manual")
    add_body(doc,
        "This section describes the features and workflows available to users registered as "
        "Service Seekers on the HomeEase platform."
    )

    # 4.3.2.1 Finding and Booking a Service Provider
    add_heading_level2(doc, "4.3.2.1 Finding and Booking a Service Provider")
    add_step(doc, 1, "After logging in as a Service Seeker, navigate to the **Find Artisans** page from the sidebar or navigation menu.")
    add_step(doc, 2, "The search page displays available service categories and a search bar. You can filter artisans by **service category**, **location**, or **rating**. The search interface is shown in Figure 4.11.")
    add_figure(doc, 11)
    add_step(doc, 3, "Enter your search criteria and click **Search** or select a category to view matching service providers.")
    add_step(doc, 4, "The search results page displays a list of verified artisans matching your criteria. Each result card shows the artisan's **name**, **service category**, **rating**, **location**, and a **Book** or **Request** button, as illustrated in Figure 4.12.")
    add_figure(doc, 12)
    add_step(doc, 5, 'Click the **Book** button on the desired artisan\'s card to initiate a service request. Fill in the required details such as **preferred date**, **time**, and a **description of the work needed**.')
    add_step(doc, 6, "Click **Submit Request** to send the booking to the service provider.")

    # 4.3.2.2 Viewing My Requests
    add_heading_level2(doc, "4.3.2.2 Viewing My Requests")
    add_step(doc, 1, "From the sidebar, click on **My Requests** to view all service requests you have submitted.")
    add_step(doc, 2, "The **My Requests** page displays a list of your requests along with their current status: **Pending**, **Accepted**, **In Progress**, **Completed**, or **Cancelled**. Each request card shows the **artisan name**, **service type**, **date**, and **status badge**, as shown in Figure 4.13.")
    add_figure(doc, 13)
    add_step(doc, 3, "Click on any request card to view the full details, including the work description and any messages exchanged with the service provider.")

    # 4.3.2.3 Making Payment
    add_heading_level2(doc, "4.3.2.3 Making Payment")
    add_step(doc, 1, "When a service request moves to the **Completed** status, a **Make Payment** option becomes available on the request detail page.")
    add_step(doc, 2, "Click the **Make Payment** button to proceed to the payment page, as shown in Figure 4.14.")
    add_figure(doc, 14)
    add_step(doc, 3, "On the payment page, review the **service details**, **amount**, and **payment method**. HomeEase uses the Flutterwave payment gateway for secure transactions.")
    add_step(doc, 4, "Select your preferred payment method (card, bank transfer, or mobile money) and click **Pay Now**.")
    add_step(doc, 5, "Follow the on-screen instructions to authorise the payment. Upon successful payment, a confirmation message is displayed and the request status is updated.")

    # 4.3.2.4 Rating a Service Provider
    add_heading_level2(doc, "4.3.2.4 Rating a Service Provider")
    add_step(doc, 1, "After completing a service and making payment, navigate to **My Requests** and locate the completed service.")
    add_step(doc, 2, "Click on the completed request to open its details. A **Rate** or **Leave a Review** option will be available, as shown in Figure 4.15.")
    add_figure(doc, 15)
    add_step(doc, 3, "Select a **star rating** (1 to 5 stars) based on the quality of service received.")
    add_step(doc, 4, "Optionally, write a **text review** describing your experience with the service provider.")
    add_step(doc, 5, "Click **Submit Review** to publish your rating. The review will be visible on the provider's profile and will contribute to their overall rating.")

    # 4.3.2.5 Messaging a Service Provider
    add_heading_level2(doc, "4.3.2.5 Messaging a Service Provider")
    add_step(doc, 1, "From the sidebar, click on **Messages** to open the messaging interface, as shown in Figure 4.16.")
    add_figure(doc, 16)
    add_step(doc, 2, "A list of your conversation threads is displayed on the left panel. Click on a thread to view the message history with a specific service provider.")
    add_step(doc, 3, "Type your message in the **message input box** at the bottom of the chat area.")
    add_step(doc, 4, "Press **Enter** or click the **Send** button to transmit your message. Messages are delivered in real time using WebSocket technology.")

    # 4.3.2.6 Managing Notifications
    add_heading_level2(doc, "4.3.2.6 Managing Notifications")
    add_step(doc, 1, "Click the **Notifications** icon (bell icon) in the navigation bar to view your latest notifications.")
    add_step(doc, 2, "Notifications include updates such as **request accepted**, **request completed**, **new message received**, and **payment confirmed**.")
    add_step(doc, 3, "Click on a notification to navigate directly to the relevant page or request detail.")
    add_step(doc, 4, "Use the **Mark All as Read** option to clear unread notification badges.")

    # 4.3.2.7 Managing Profile
    add_heading_level2(doc, "4.3.2.7 Managing Profile")
    add_step(doc, 1, "From the sidebar, click on **Profile** or **Account Settings** to access your profile management page.")
    add_step(doc, 2, "You can update your **Full Name**, **Email Address**, **Phone Number**, and **Profile Picture**.")
    add_step(doc, 3, "Click **Save Changes** to update your profile information.")
    add_step(doc, 4, "To change your password, navigate to the **Change Password** section, enter your **current password** and **new password**, then click **Update Password**.")

    # ── 4.3.3 Service Provider User Manual ──
    add_heading_level2(doc, "4.3.3 Service Provider User Manual")
    add_body(doc,
        "This section describes the features and workflows available to users registered as "
        "Service Providers on the HomeEase platform."
    )

    # 4.3.3.1 Awaiting Verification
    add_heading_level2(doc, "4.3.3.1 Awaiting Verification")
    add_step(doc, 1, "After completing the registration form and submitting your details, your account is placed in a **Pending Verification** state.")
    add_step(doc, 2, "Upon logging in, you will see a verification pending screen informing you that your account is awaiting administrator review, as shown in Figure 4.17.")
    add_figure(doc, 17)
    add_step(doc, 3, "During this period, you cannot accept job offers or access provider-specific features.")
    add_step(doc, 4, "Once an administrator verifies your account, you will receive a notification and gain full access to the provider dashboard.")

    # 4.3.3.2 Chatting with Admin During Verification
    add_heading_level2(doc, "4.3.3.2 Chatting with Admin During Verification")
    add_step(doc, 1, "While your account is pending verification, a **Chat with Admin** option is available on the verification pending screen.")
    add_step(doc, 2, "Click on **Chat with Admin** to open a direct messaging thread with the platform administrator, as shown in Figure 4.18.")
    add_figure(doc, 18)
    add_step(doc, 3, "Use this channel to submit any additional documents, answer verification queries, or seek clarification about the onboarding process.")
    add_step(doc, 4, "Messages are delivered in real time, and the admin can respond directly within the same chat window.")

    # 4.3.3.3 Viewing the Provider Dashboard
    add_heading_level2(doc, "4.3.3.3 Viewing the Provider Dashboard")
    add_body(doc,
        "Once verified, the provider dashboard becomes accessible upon login. The dashboard "
        "provides an overview of key metrics including total earnings, active jobs, completed "
        "jobs, and pending requests. The dashboard also serves as the central hub for "
        "navigating to other provider features such as **My Jobs**, **Transactions**, and "
        "**Profile** management."
    )

    # 4.3.3.4 Managing Job Offers
    add_heading_level2(doc, "4.3.3.4 Managing Job Offers")
    add_step(doc, 1, "Navigate to **My Jobs** from the sidebar to view all job offers received from service seekers.")
    add_step(doc, 2, "Each job offer displays the **seeker's name**, **service requested**, **preferred date and time**, **location**, and **job description**.")
    add_step(doc, 3, "To accept a job, click the **Accept** button on the offer card. The seeker will be notified of your acceptance.")
    add_step(doc, 4, "To decline a job, click the **Decline** button. Optionally, provide a reason for declining.")
    add_step(doc, 5, "Accepted jobs move to the **Active Jobs** section where you can manage their progress.")

    # 4.3.3.5 Checking In and Checking Out
    add_heading_level2(doc, "4.3.3.5 Checking In and Checking Out")
    add_step(doc, 1, "On the day of a scheduled job, navigate to **My Jobs** and locate the active job.")
    add_step(doc, 2, "Click the **Check In** button when you arrive at the service location. This records your start time and notifies the service seeker that you have arrived, as shown in Figure 4.19.")
    add_figure(doc, 19)
    add_step(doc, 3, "Upon completing the work, click the **Check Out** button to record your end time.")
    add_step(doc, 4, "The job status is automatically updated to **Completed**, and the service seeker is notified to review the service and make payment.")

    # 4.3.3.6 Viewing Transactions and Earnings
    add_heading_level2(doc, "4.3.3.6 Viewing Transactions and Earnings")
    add_step(doc, 1, "From the sidebar, click on **Transactions** or **Earnings** to view your financial history.")
    add_step(doc, 2, "The transactions page displays a list of all payments received, including the **job description**, **amount**, **date**, and **status** (completed, pending, or processed).")
    add_step(doc, 3, "A summary of total earnings, pending payouts, and processed payouts is displayed at the top of the page.")
    add_step(doc, 4, "Click on any transaction to view additional details such as the seeker's name and the service category.")

    # 4.3.3.7 Managing Profile
    add_heading_level2(doc, "4.3.3.7 Managing Profile")
    add_step(doc, 1, "Navigate to **Profile** from the sidebar to manage your provider profile.")
    add_step(doc, 2, "Update your **personal information**, **skills and services**, **portfolio images**, and **bank details** as needed.")
    add_step(doc, 3, "Adding a professional **profile picture** and detailed **service descriptions** can help attract more service seekers.")
    add_step(doc, 4, "Click **Save Changes** to persist any updates to your profile.")

    # ── 4.3.4 Administrator User Manual ──
    add_heading_level2(doc, "4.3.4 Administrator User Manual")
    add_body(doc,
        "This section describes the administrative features available to platform administrators "
        "for managing users, verifying providers, processing payouts, and handling support requests."
    )

    # 4.3.4.1 Accessing the Admin Dashboard
    add_heading_level2(doc, "4.3.4.1 Accessing the Admin Dashboard")
    add_step(doc, 1, "Navigate to the HomeEase login page and enter the administrator **Email Address** and **Password**.")
    add_step(doc, 2, "Upon successful login, the system detects the admin role and redirects to the **Admin Dashboard**.")
    add_step(doc, 3, "The admin dashboard displays an overview of key platform metrics: **total users**, **total providers**, **total service seekers**, **active jobs**, **completed transactions**, and **revenue summary**.")
    add_step(doc, 4, "The sidebar provides navigation links to all administrative sections: **Users**, **Providers**, **Payouts**, **Messages**, and **Settings**.")

    # 4.3.4.2 Verifying Service Providers
    add_heading_level2(doc, "4.3.4.2 Verifying Service Providers")
    add_step(doc, 1, "From the admin sidebar, click on **Providers** to view the list of all registered service providers.")
    add_step(doc, 2, "Filter the list by status to show only **Pending Verification** providers.")
    add_step(doc, 3, "Click on a pending provider to review their registration details, including **personal information**, **skills**, **uploaded documents**, and **bank details**.")
    add_step(doc, 4, "If the provider meets the platform's verification criteria, click the **Verify** or **Approve** button to activate their account.")
    add_step(doc, 5, "If additional information is required, use the **Chat** feature to communicate with the provider directly.")
    add_step(doc, 6, "To reject a provider application, click **Reject** and optionally provide a reason.")

    # 4.3.4.3 Processing Payouts
    add_heading_level2(doc, "4.3.4.3 Processing Payouts")
    add_step(doc, 1, "From the admin sidebar, click on **Payouts** to access the payout management screen, as shown in Figure 4.20.")
    add_figure(doc, 20)
    add_step(doc, 2, "The payouts screen displays a list of all pending payout requests from service providers. Each entry shows the **provider name**, **amount**, **bank details**, **job reference**, and **request date**.")
    add_step(doc, 3, "Review each payout request to ensure the service has been completed and the amount is correct.")
    add_step(doc, 4, 'Click the **Confirm** or **Process Payout** button next to a provider\'s entry to initiate the payout, as shown in Figure 4.21.')
    add_figure(doc, 21)
    add_step(doc, 5, "A confirmation dialog appears. Verify the details once more and click **Confirm Payout** to finalise the transaction.")
    add_step(doc, 6, "The provider will receive a notification that their payout has been processed, and the payout status will be updated to **Completed**.")

    # 4.3.4.4 Managing Users and Viewing Statistics
    add_heading_level2(doc, "4.3.4.4 Managing Users and Viewing Statistics")
    add_step(doc, 1, "Click on **Users** in the admin sidebar to view a comprehensive list of all registered users on the platform.")
    add_step(doc, 2, "Use the search and filter options to find specific users by **name**, **email**, **role**, or **registration date**.")
    add_step(doc, 3, "Click on a user to view their full profile, activity history, and transaction records.")
    add_step(doc, 4, "The **Statistics** section of the dashboard provides visual charts and graphs showing platform growth, service demand trends, and revenue analytics.")
    add_step(doc, 5, "Administrators can **suspend** or **deactivate** user accounts if violations of platform policies are detected.")

    # 4.3.4.5 Handling Support Messages
    add_heading_level2(doc, "4.3.4.5 Handling Support Messages")
    add_step(doc, 1, "From the admin sidebar, click on **Messages** or **Support** to view all incoming support messages.")
    add_step(doc, 2, "Messages are organised by user and displayed in a chat-like interface, similar to the messaging feature available to seekers and providers.")
    add_step(doc, 3, "Click on a conversation thread to read the message history and respond to the user.")
    add_step(doc, 4, "Type your response in the message input box and press **Enter** or click **Send** to deliver the message.")
    add_step(doc, 5, "Common support queries include account verification status, payment issues, and platform usage guidance.")
    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.4 System Architecture
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.4 System Architecture")
    add_body(doc,
        "The HomeEase platform follows a client-server architecture built on modern web "
        "technologies. The frontend is developed using React.js with Tailwind CSS for "
        "responsive and visually appealing user interfaces. The backend is powered by Node.js "
        "with the Express.js framework, providing RESTful API endpoints for all platform "
        "operations. Data persistence is handled by MongoDB, a NoSQL database that offers "
        "flexibility and scalability for the platform's data model. Real-time communication "
        "between clients and the server is implemented using Socket.IO, enabling instant "
        "messaging and live notifications. The Flutterwave payment gateway is integrated for "
        "secure and reliable payment processing. Cloud-based image storage is managed through "
        "Cloudinary. Figure 4.22 illustrates the complete system architecture diagram."
    )
    add_figure(doc, 22)
    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.5 Sample System Run
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.5 Sample System Run")
    add_body(doc,
        "This section presents a complete end-to-end walkthrough of a typical service booking "
        "scenario on the HomeEase platform. The narrative follows a service seeker named Ada "
        "as she searches for, books, and pays for a plumbing service from a verified provider "
        "named Emeka."
    )

    steps_narrative = [
        ("Ada opens her web browser and navigates to the HomeEase platform URL. The landing page loads, displaying the hero section, services grid, and feature highlights.",
         "The landing page (Figure 4.1) provides an inviting introduction to the platform, showcasing the range of available services."),

        ("Ada scrolls through the landing page to review the services offered. She sees categories such as Plumbing, Electrical, Cleaning, Painting, and more (Figure 4.2).",
         "The services grid gives Ada confidence that the platform covers the specific service she needs."),

        ("Having identified that she needs a plumber, Ada clicks the **Sign Up** button. She is presented with the role selection screen (Figure 4.5) and selects **Service Seeker**."),

        ("Ada fills in her registration details — **Full Name**: Ada Okonkwo, **Email**: ada@example.com, **Phone Number**: 08012345678, and **Password** — on the Service Seeker registration form (Figure 4.6). She clicks **Create Account**."),

        ("After successful registration, Ada is redirected to the login page. She enters her email and password (Figure 4.10) and clicks **Login** to access her seeker dashboard."),

        ("From the sidebar navigation, Ada clicks on **Find Artisans**. The search page loads (Figure 4.11), displaying service categories and a search bar."),

        ("Ada selects **Plumbing** from the service categories or types 'Plumber' into the search bar. She further narrows the search by selecting her location and clicks **Search**."),

        ("The search results page displays a list of verified plumbers in her area (Figure 4.12). Ada reviews the profiles, noting each plumber's rating, location, and number of completed jobs."),

        ("Impressed by his 4.8-star rating and proximity, Ada clicks the **Book** button on Emeka's profile card. A booking form appears."),

        ("Ada fills in the booking details: **Service**: Plumbing (Pipe Repair), **Preferred Date**: 15th March 2025, **Preferred Time**: 10:00 AM, and **Description**: 'Kitchen sink pipe is leaking. Need urgent repair.' She clicks **Submit Request**."),

        ("Ada navigates to **My Requests** (Figure 4.13) to check the status of her request. The request shows as **Pending**, awaiting Emeka's response."),

        ("Emeka, the plumber, receives a notification about Ada's booking request. He reviews the details, accepts the job, and sends a message via the chat system confirming the appointment."),

        ("Ada receives a notification that her request has been accepted. She opens the **Messages** section (Figure 4.16) and sees Emeka's confirmation message. She replies to confirm the time and address."),

        ("On the scheduled day, Emeka arrives at Ada's location. He opens his **My Jobs** page (Figure 4.19) and clicks **Check In** to record his arrival time. Ada receives a real-time notification."),

        ("Emeka completes the pipe repair work. He then clicks **Check Out** on the My Jobs page to mark the job as completed."),

        ("Ada receives a notification that the service has been completed. She opens **My Requests** and sees the request status has changed to **Completed** (Figure 4.15)."),

        ("Ada clicks on the completed request and selects the option to **Rate** the service. She gives Emeka a 5-star rating and writes a positive review: 'Excellent work! Very professional and timely.' She clicks **Submit Review** (Figure 4.15)."),

        ("Ada then clicks the **Make Payment** button on the request detail page. The payment page loads (Figure 4.14), showing the service details and amount of ₦15,000."),

        ("Ada selects her preferred payment method (debit card) and clicks **Pay Now**. She is redirected to the Flutterwave secure payment gateway where she enters her card details and authorises the payment."),

        ("The payment is processed successfully. Ada sees a confirmation message, and the request status in **My Requests** updates to **Paid**."),

        ("On the administrative side, the platform administrator logs into the admin dashboard. Under **Payouts**, Emeka's earnings of ₦15,000 (minus platform commission) appear as a pending payout (Figure 4.20). The admin reviews the details and clicks **Confirm** to process the payout (Figure 4.21). Emeka receives a notification that his payout has been initiated."),

        ("The entire transaction is recorded in the system. Ada can view the completed service in her history, Emeka's profile now reflects the new 5-star review, and the platform's analytics update to include this completed transaction."
        ),
    ]

    for i, (step_text, *_) in enumerate(steps_narrative, 1):
        add_step(doc, i, step_text)

    add_page_break(doc)

    # ═══════════════════════════════════════════
    # 4.6 Summary
    # ═══════════════════════════════════════════
    add_heading_level1(doc, "4.6 Summary")
    add_body(doc,
        "This chapter has provided a thorough walkthrough of the HomeEase platform's system "
        "implementation and user guide. The implementation environment, including hardware and "
        "software requirements, was documented to ensure reproducibility. Detailed step-by-step "
        "instructions were provided for all three user roles — Service Seekers, Service "
        "Providers, and Administrators — covering every major feature from registration and "
        "login to service booking, payment processing, ratings and reviews, real-time "
        "messaging, provider verification, and payout management. Twenty-two annotated "
        "screenshots were included throughout the chapter to illustrate the system's interfaces "
        "and workflows. The system architecture was presented, describing the technologies and "
        "design patterns that underpin the platform. Finally, a 22-step sample system run "
        "demonstrated a complete end-to-end service booking scenario, validating the system's "
        "functionality and usability. The HomeEase platform, as demonstrated in this chapter, "
        "provides a robust, user-friendly, and comprehensive solution for connecting service "
        "seekers with verified service providers in Nigeria."
    )

    # ── Save ──
    output_path = os.path.join(BASE, "HomeEase_Chapter_4_Walkthrough.docx")
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"Document saved to: {output_path}")
    print(f"File size: {size_kb:.1f} KB")
    if size_kb < 500:
        print("WARNING: File is smaller than 500 KB — screenshots may be missing!")
    else:
        print("OK: File size exceeds 500 KB.")


if __name__ == "__main__":
    main()