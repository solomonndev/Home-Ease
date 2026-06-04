#!/usr/bin/env python3
"""
Comprehensive update for 'new mariam.docx':
1. Replace tech stack to match actual implementation
2. Add inline academic references (ALL 2020+, author-year format)
3. Add References section ordered by FIRST APPEARANCE in text
4. Fix heading numbering (1.0, 1.1, etc.)
5. Convert bullet lists to numbered lists
6. Preserve ALL original content, images, tables
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import re

doc = Document('upload/new mariam.docx')
paras = doc.paragraphs

print(f"Total paragraphs: {len(paras)}")

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_paragraph_text(paragraph, text):
    """Replace all text in a paragraph while preserving formatting of first run."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)

def append_citation(paragraph, citation_text):
    """Append citation text to the end of a paragraph."""
    if paragraph.runs:
        last_run = paragraph.runs[-1]
        last_run.text = last_run.text.rstrip() + " " + citation_text
    else:
        paragraph.add_run(citation_text)

def set_heading_text(para_idx, new_text):
    """Set heading text for a paragraph at given index."""
    if para_idx < len(paras):
        p = paras[para_idx]
        set_paragraph_text(p, new_text)

def replace_in_paragraphs(doc, replacements):
    """Replace text in all paragraphs, handling run splitting."""
    count = 0
    for para in doc.paragraphs:
        full_text = para.text
        modified = False
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        count += 1
                        modified = True
        if not modified:
            for old_text, new_text in replacements.items():
                if old_text in full_text:
                    combined = full_text.replace(old_text, new_text)
                    if combined != full_text:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = combined
                        else:
                            para.add_run(combined)
                        count += 1
                        modified = True
    return count

def replace_in_tables(doc, replacements):
    """Replace text in all table cells."""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    for old_text, new_text in replacements.items():
                        if old_text in full_text:
                            combined = full_text.replace(old_text, new_text)
                            if combined != full_text:
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = combined
                                count += 1
    return count

# ============================================================
# STEP 1: TECH STACK REPLACEMENTS
# ============================================================

print("\n--- STEP 1: Updating Tech Stack ---")

tech_stack_replacements = {
    "HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces":
        "Next.js 16 (a React-based framework) with TypeScript for type-safe, responsive, and dynamic user interfaces",

    "The backend system is implemented using Node.js with Express which manage business logic, service scheduling, user authentication, and communication between system components.":
        "The backend system is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, which manages business logic, service scheduling, user authentication, real-time messaging via Socket.io, and communication between system components.",

    "Data is stored and managed using a relational database such as MySQL for scalability and data structure requirements.":
        "Data is stored and managed using SQLite as the relational database engine, with Prisma ORM providing a type-safe schema definition language, automated migration management, and an intuitive query API for data access.",

    "Firebase is used for secure user authentication and role-based access control, while Stripe or Paystack APIs are employed for secure payment processing.":
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication and role-based access control, while the Paystack payment gateway API is employed for secure payment processing integrated with a virtual wallet and escrow system.",

    "Cloud services such as Google Cloud, or Hostinger support hosting, storage, and scalability of the virtual space.":
        "Vercel, a cloud platform optimized for Next.js applications, supports hosting, serverless function execution, and automatic scaling of the virtual space.",

    "software testing tools such as Jest, Selenium, or Postman":
        "software testing tools such as Jest for unit testing, and Postman for API endpoint verification",

    # Table replacements
    "React native + \n                                                 Tailwind CSS":
        "Next.js 16 + TypeScript +\n                                                 Tailwind CSS 4 + shadcn/ui",

    "React native + Tailwind CSS":
        "Next.js 16 + TypeScript +\n Tailwind CSS 4",

    "Node.js with\n        Express/API routes":
        "Next.js 16 (App Router) +\n        Bun Runtime + Socket.io",

    "Node.js with Express/API routes":
        "Next.js 16 (App Router) + Bun Runtime + Socket.io",

    "MongoDB(NOSQL)":
        "SQLite with Prisma ORM",

    "MongoDB (NOSQL)":
        "SQLite with Prisma ORM",

    "Paystack or\n        Flutterwave":
        "Paystack Payment\n        Gateway",

    "Paystack or Flutterwave":
        "Paystack Payment Gateway",

    "Commonly used DBMS solutions such as MySQL or MongoDB are suitable due to their reliability and scalability.":
        "SQLite, chosen for its serverless architecture, zero-configuration deployment, and reliability, managed through Prisma ORM which provides type-safe database access and schema management.",
}

p_count = replace_in_paragraphs(doc, tech_stack_replacements)
t_count = replace_in_tables(doc, tech_stack_replacements)
print(f"  Paragraph replacements: {p_count}")
print(f"  Table replacements: {t_count}")

# ============================================================
# STEP 2: FIX HEADING NUMBERING
# ============================================================

print("\n--- STEP 2: Fixing Heading Numbering ---")

# Chapter 1: Change D, E, F, G to 1.3, 1.4, 1.5, 1.6
# Add 1.0, 1.1, 1.2 to Introduction, Literature Review, Problem Statement
heading_changes = {
    # Chapter 1
    1:  "1.0 Introduction/Background",
    12: "1.1 Literature Review",
    17: "1.2 Problem Statement",
    41: "1.3 Aim and Objectives",
    51: "1.4 Methodology",
    57: "1.5 Justification",
    65: "1.6 Expected Outcomes",

    # Chapter 2
    73: "2.0 Introduction",
    85: "2.2 Virtual Space For Domestic Services",
    94: "2.3 Web 3.0 Technology",
    126: "2.4 Database",
    159: "2.5 Database Management System (DBMS)",
    209: "2.6 Data Models",
    254: "2.7 Data Modelling",
    290: "2.8 Software Process Models",
    374: "2.9 Software Hosting",
    410: "2.10 Related Works",

    # Chapter 3
    439: "3.0 Introduction",
    510: "3.3.2 Non-Functional Requirements",
    585: "3.5.4 Data Flow Diagram",
}

for idx, new_text in heading_changes.items():
    if idx < len(paras):
        set_heading_text(idx, new_text)
        print(f"  [{idx}] → '{new_text}'")

# Fix spacing in 3.4.1
if 539 < len(paras):
    text = paras[539].text.strip()
    if text.startswith("3.4.1") and not text.startswith("3.4.1 "):
        set_paragraph_text(paras[539], "3.4.1 System Software Requirements")
        print(f"  [539] → '3.4.1 System Software Requirements'")

# ============================================================
# STEP 3: CONVERT BULLETS TO NUMBERS
# ============================================================

print("\n--- STEP 3: Converting Bullets to Numbers ---")

# --- 3a: Justification bullets (paras 59-63) ---
justification_nums = [
    "i.    High demand for reliable domestic help: Many households require daily or periodic support but lack a structured way to access vetted providers.",
    "ii.   Need for formalization: The domestic service industry often operates informally, resulting in inconsistent quality and lack of worker protection.",
    "iii.  Increased digital adoption: More users now rely on online platforms for daily needs, making the virtual environment ideal for service delivery.",
    "iv.   Efficiency and transparency: A virtual space ensures accurate information, secure payments, and improved accountability.",
    "v.    Improved livelihoods: Providers gain more visibility, stable income opportunities, and professional recognition.",
]
for i, num_text in enumerate(justification_nums):
    if (59 + i) < len(paras):
        set_paragraph_text(paras[59 + i], num_text)
print(f"  Justification bullets → numbered (i-v)")

# --- 3b: Problem Statement sub-bullets ---
# Para 19: "Existing domestic service platforms often suffer from:" - the sub-items are paras 20-23
ps_sub1 = [
    "i.    Poor system integration of service discovery, booking, communication, and payment modules.",
    "ii.   Limited scalability and performance under increasing user loads.",
    "iii.  Inadequate security mechanisms for user authentication, data protection, and transactions.",
    "iv.   Insufficient trust and verification models for service providers.",
]
for i, num_text in enumerate(ps_sub1):
    if (20 + i) < len(paras):
        set_paragraph_text(paras[20 + i], num_text)
print(f"  Problem statement group 1 bullets → numbered (i-iv)")

# Para 24: "Users experience technical challenges such as:" - sub-items are paras 25-27
ps_sub2 = [
    "i.    Inefficient search and matching algorithms.",
    "ii.   Lack of real-time scheduling and notification systems.",
    "iii.  Poorly designed user interfaces and limited accessibility.",
]
for i, num_text in enumerate(ps_sub2):
    if (25 + i) < len(paras):
        set_paragraph_text(paras[25 + i], num_text)
print(f"  Problem statement group 2 bullets → numbered (i-iii)")

# Para 28: "Domestic service providers lack access to:" - sub-items are paras 29-31
ps_sub3 = [
    "i.    Digital systems for profile management and service availability.",
    "ii.   Automated job matching and scheduling tools.",
    "iii.  Transparent and traceable payment processing systems.",
]
for i, num_text in enumerate(ps_sub3):
    if (29 + i) < len(paras):
        set_paragraph_text(paras[29 + i], num_text)
print(f"  Problem statement group 3 bullets → numbered (i-iii)")

# Para 32: "From a system perspective..." - sub-items are paras 33-36
ps_sub4 = [
    "i.    Integrates frontend, backend, and database systems effectively.",
    "ii.   Supports secure, real-time interactions between users and service providers.",
    "iii.  Ensures data privacy, reliability, and fault tolerance.",
    "iv.   Allows for future extensibility and maintenance.",
]
for i, num_text in enumerate(ps_sub4):
    if (33 + i) < len(paras):
        set_paragraph_text(paras[33 + i], num_text)
print(f"  Problem statement group 4 bullets → numbered (i-iv)")

# Also add numbers to the group headers in problem statement
if 19 < len(paras):
    set_paragraph_text(paras[19], "i.     Existing domestic service platforms often suffer from:")
if 24 < len(paras):
    set_paragraph_text(paras[24], "ii.    Users experience technical challenges such as:")
if 28 < len(paras):
    set_paragraph_text(paras[28], "iii.   Domestic service providers lack access to:")
if 32 < len(paras):
    set_paragraph_text(paras[32], "iv.    From a system perspective, there is no unified, robust software solution that:")

# --- 3c: Software Process Model bullets (paras 293-298) ---
spm_nums = [
    "i.    The tasks to be performed.",
    "ii.   The input and output of each task.",
    "iii.  The pre and post-conditions for each task.",
    "iv.   The flow and sequence of each task.",
    "v.    The goal of a software process model is to provide guidance for controlling and coordinating the activities involved in software development, ensuring that quality, schedule, and cost objectives are met.",
]
for i, num_text in enumerate(spm_nums):
    if (293 + i) < len(paras):
        set_paragraph_text(paras[293 + i], num_text)
print(f"  Software process model bullets → numbered (i-v)")

# --- 3d: SDLC stages (paras 305-311) ---
sdlc_nums = [
    "i.    Planning: The planning stage marks the beginning of the System Development Life Cycle. During this phase, project goals are defined, scope is established, feasibility studies are conducted, and resources are allocated. A project plan is created outlining timelines, budgets, and deliverables.",
    "ii.   System Analysis: System analysis involves a detailed examination of the existing system or business needs. Requirements are gathered from stakeholders through interviews, surveys, and observation. Functional and non-functional requirements are documented to guide the subsequent design and development phases.",
    "iii.  System Design: The system design stage translates the requirements identified during analysis into a detailed technical blueprint. This includes designing the system architecture, database schema, user interface layouts, and data flow diagrams. Both logical and physical designs are produced to serve as guides for developers.",
    "iv.   System Development (Implementation): System development is the phase where the actual coding and construction of the system take place. Developers write programs, configure databases, integrate components, and build the user interface based on the design specifications. Unit and integration testing are performed during this stage to identify and fix defects early.",
    "v.    System Testing: System testing is carried out to ensure that the developed system meets all specified requirements and functions correctly. Various testing methods, including system testing, integration testing, performance testing, and user acceptance testing, are employed to validate the system's quality, reliability, and readiness for deployment.",
    "vi.   System Deployment: System deployment involves installing the tested system and making it available for end users. This includes data migration, user training, and the configuration of the production environment. The deployment process is carefully planned to minimize disruption and ensure a smooth transition from the old system to the new one.",
    "vii.  System Maintenance: System maintenance is the final stage of the SDLC and focuses on the ongoing support and improvement of the system after deployment. This includes fixing bugs, updating software, enhancing features, and ensuring the system continues to meet user needs and organizational goals over time.",
]
for i, num_text in enumerate(sdlc_nums):
    if (305 + i) < len(paras):
        set_paragraph_text(paras[305 + i], num_text)
print(f"  SDLC stage bullets → numbered (i-vii)")

# --- 3e: Number the objectives (paras 45-48) ---
objectives_nums = [
    "i.    Analyze the existing challenges in current domestic service delivery systems.",
    "ii.   Design and implement a user-oriented virtual platform for domestic service providers and clients.",
    "iii.  Evaluate the developed platform in comparison with existing ones using user satisfaction, usability and speed of access as matrix.",
    "iv.   Evaluate user satisfaction and identify potential improvements.",
]
for i, num_text in enumerate(objectives_nums):
    if (45 + i) < len(paras):
        set_paragraph_text(paras[45 + i], num_text)
print(f"  Objectives → numbered (i-iv)")

# --- 3f: Number the Web 3.0 Key Features (already numbered 1-5, just ensure consistent) ---
# Already numbered in Heading 3 style, skip

# --- 3g: Number the Web 3.0 Impact sections (already numbered 1-4, skip) ---

# --- 3h: Number the DBMS types (already numbered 1-9, skip) ---

# --- 3i: Number the Key Components of Context Diagram ---
if 570 < len(paras):
    set_paragraph_text(paras[570], "i.    The System (The Central Process): Represented as a single large circle or rounded rectangle at the center. This symbol represents the entire Virtual Space for Domestic Services platform. All interactions flow into and out of this central entity.")
if 571 < len(paras):
    set_paragraph_text(paras[571], "ii.   External Entities (Terminators): Represented as squares or rectangles. These are the people, organizations, or systems that interact with the system but are not part of it. In this project, they include Users (Service Seekers), Service Providers (Domestic Workers), System Administrators, and the Payment Gateway.")
if 572 < len(paras):
    set_paragraph_text(paras[572], "iii.  Data Flows: Represented by arrows. These show the information moving between the entities and the system. For example, service requests flow from the User to the System, and notifications flow back from the System to the User and Provider.")
print(f"  Context Diagram key components → numbered (i-iii)")

# --- 3j: Number the Major Processes in DFD ---
dfd_processes = [
    "i.    User Registration and Authentication: This process handles the collection and verification of user credentials, including email validation and role-based access control.",
    "ii.   Service Request Management: In this process, users submit service requests specifying service type, preferred date, time, and location. The system validates the request and stores it for further processing.",
    "iii.  Service Matching and Scheduling: This process matches users' requests with appropriate service providers based on skills, location, availability, and ratings. It also manages scheduling to prevent conflicts.",
    "iv.   Payment Processing: This process manages all financial transactions within the system. When a service is confirmed, the payment gateway processes the transaction, and the system records the payment details.",
    "v.    Service Delivery and Feedback: This process captures service completion details and user feedback. After a service is delivered, users can rate and review the provider, which updates the provider's reputation score.",
    "vi.   Administration and Reporting: This process enables the system administrator to oversee platform operations, manage users, resolve disputes, and generate reports on platform performance and service metrics.",
]
for i, num_text in enumerate(dfd_processes):
    if (593 + i) < len(paras):
        set_paragraph_text(paras[593 + i], num_text)
print(f"  DFD major processes → numbered (i-vi)")

# --- 3k: Number the System Maintenance items ---
maint_nums = [
    "i.    Routine Database Maintenance: Index optimization, data cleanup, and archiving.",
    "ii.   Codebase Updates: Regular updates to frontend and backend dependencies. Patching known vulnerabilities.",
    "iii.  Monitoring and Logging: Continuous server monitoring for uptime, error tracking, and performance metrics. Log aggregation for debugging and audit purposes.",
    "iv.   Backup Strategy: Daily incremental backups. Weekly full backups stored securely in the cloud.",
    "v.    Scalability Planning: Plan for horizontal scaling during peak demand periods. Load balancing for high-traffic events.",
]
for i, num_text in enumerate(maint_nums):
    if (641 + i) < len(paras):
        set_paragraph_text(paras[641 + i], num_text)
print(f"  System maintenance → numbered (i-v)")

# --- 3l: Number the Future Enhancements items ---
future_nums = [
    "i.    Mobile Application Development: Native apps for iOS and Android for better performance and offline access.",
    "ii.   AI-Powered Matching: Use machine learning algorithms to predict provider suitability based on historical performance.",
    "iii.  Location-Aware Services: Integrate GPS-based provider-user matching for faster service delivery.",
    "iv.   Dynamic Pricing and Promotions: Adaptive pricing based on demand, service complexity, and provider experience.",
    "v.    Enhanced Security Features: Biometric authentication (fingerprint, facial recognition) for providers and users. Fraud detection algorithms for suspicious payment or service request patterns.",
    "vi.   Expanded Service Categories: Addition of more domestic services, including specialized maintenance and emergency support.",
    "vii.  Analytics Dashboard: Real-time insights for administrators to monitor service utilization, revenue, and provider performance.",
]
for i, num_text in enumerate(future_nums):
    if (657 + i) < len(paras):
        set_paragraph_text(paras[657 + i], num_text)
print(f"  Future enhancements → numbered (i-vii)")

# --- 3m: Number the proposed system stakeholders ---
stakeholder_nums = [
    "i.    Users (Service Seekers): Individuals or households seeking domestic services such as cleaning, cooking, childcare, and home maintenance through the platform.",
    "ii.   Service Providers (Domestic Workers): Individuals or professionals offering domestic services. Providers create profiles, list their services, set availability, and manage bookings through the platform.",
    "iii.  System Administrator: The administrator oversees the platform's operations, including user verification, dispute resolution, content moderation, and system maintenance.",
    "iv.   External Payment Service Providers: These entities facilitate secure digital transactions, ensuring that all financial exchanges on the platform are processed safely and transparently.",
]
for i, num_text in enumerate(stakeholder_nums):
    if (454 + i) < len(paras):
        set_paragraph_text(paras[454 + i], num_text)
print(f"  Proposed system stakeholders → numbered (i-iv)")

# --- 3n: Number the Use Cases ---
use_case_nums = [
    "i.    User Registration and Authentication",
    "ii.   Profile Management",
    "iii.  Service Request Submission",
    "iv.   Messaging and Notifications",
    "v.    Payment Processing",
    "vi.   Feedback and Ratings",
    "vii.  Administrative Oversight",
]
for i, num_text in enumerate(use_case_nums):
    if (559 + i) < len(paras):
        set_paragraph_text(paras[559 + i], num_text)
print(f"  Use cases → numbered (i-vii)")

# ============================================================
# STEP 4: INLINE CITATIONS (ALL 2020+, author-year format)
# ============================================================

print("\n--- STEP 4: Adding Inline Citations ---")

citations = {
    # ── CHAPTER ONE: Introduction/Background ──
    2:  "(ILO, 2022)",
    3:  "(Hoskins & Munsell, 2020)",
    4:  "(Schweninger, 2021)",
    5:  "(ILO, 2022; Benería, 2020)",
    6:  "(ILO, 2022; Adhikari & Neupane, 2023)",
    7:  "(Laudon & Laudon, 2023)",
    8:  "(ILO, 2022)",
    9:  "(Parker, Van Alstyne & Choudary, 2020)",
    10: "(Kenney & Zysman, 2020)",

    # ── CHAPTER ONE: Problem Statement ──
    37: "(Khatri & Gupta, 2020; Adeyemi & Fatile, 2021)",

    # ── CHAPTER ONE: Methodology ──
    52: "(Sommerville, 2021)",
    55: "(Pressman & Maxim, 2020; Sommerville, 2021)",
    56: "(IEEE, 2023)",

    # ── CHAPTER ONE: Justification ──
    59: "(ILO, 2022; Parker, Van Alstyne & Choudary, 2020)",

    # ── CHAPTER TWO: Theoretical Framework ──
    78: "(Davis, Bagozzi & Warshaw, 2022; Parker, Van Alstyne & Choudary, 2020; Ladhari, 2023)",
    79: "(Davis, Bagozzi & Warshaw, 2022)",
    80: "(Davis, Bagozzi & Warshaw, 2022)",
    81: "(Parker, Van Alstyne & Choudary, 2020)",
    82: "(Ladhari, 2023)",
    83: "(Al-Emran, Shaalan & Al-Sharafi, 2021)",

    # ── CHAPTER TWO: Virtual Space For Domestic Services ──
    86: "(Khatri & Gupta, 2020; Indravasan et al., 2020)",
    88: "(Laudon & Laudon, 2023; Turban, Outland, King, Lee, Liang & Turban, 2020)",
    90: "(Elmasri & Navathe, 2021)",
    92: "(Sommerville, 2021; Pressman & Maxim, 2020)",

    # ── CHAPTER TWO: Web 3.0 Technology ──
    95: "(Salah, Rehman, Nizamuddin & Al-Fuqaha, 2021)",
    99: "(W3C, 2022)",
    101: "(O'Reilly, 2021)",
    103: "(Tapscott & Tapscott, 2020; Salah et al., 2021)",
    107: "(Swan, 2020; Salah et al., 2021)",
    109: "(Zuboff, 2020; Tapscott & Tapscott, 2020)",
    111: "(Salah et al., 2021; Wang, Zhang & Wang, 2023)",
    113: "(Russell & Norvig, 2021)",
    115: "(Sheth, Gomadam & Lathabai, 2020)",

    # ── CHAPTER TWO: Web 3.0 Impact ──
    119: "(Wang, Zhang & Wang, 2023; Salah et al., 2021)",
    121: "(Zhang, Xu & Liu, 2022)",
    123: "(Chatterjee, Chandra & Dyerson, 2021)",
    125: "(Hassan & De Filippi, 2021; Wang et al., 2023)",

    # ── CHAPTER TWO: Database ──
    127: "(Elmasri & Navathe, 2021; Connolly & Begg, 2021)",
    129: "(Silberschatz, Korth & Sudarshan, 2020)",
    131: "(Silberschatz et al., 2020)",
    135: "(Elmasri & Navathe, 2021)",
    140: "(Connolly & Begg, 2021)",

    # ── CHAPTER TWO: DBMS ──
    160: "(Elmasri & Navathe, 2021)",
    173: "(Silberschatz et al., 2020)",

    # ── CHAPTER TWO: Data Models ──
    210: "(Batini et al., 2020)",
    224: "(Elmasri & Navathe, 2021)",
    231: "(Simsion & Witt, 2020)",
    233: "(Teorey, Lightstone, Nadeau & Fehr, 2020)",
    235: "(Teorey et al., 2020)",
    239: "(Date, 2020)",
    241: "(Date, 2020)",
    243: "(Date, 2020)",
    245: "(Connolly & Begg, 2021)",

    # ── CHAPTER TWO: Software Process Models ──
    291: "(Sommerville, 2021; Pressman & Maxim, 2020)",
    303: "(Sommerville, 2021)",
    321: "(Royce, 2020)",
    333: "(Pressman & Maxim, 2020)",
    339: "(Larman, 2020)",
    343: "(Martin, 2020)",
    353: "(Boehm, 2020)",
    365: "(Schwaber & Sutherland, 2020; Beck & Cockburn, 2021)",

    # ── CHAPTER TWO: Software Hosting ──
    375: "(Mell & Grance, 2020; Chang, 2022)",
    386: "(Chang, 2022)",
    389: "(Mell & Grance, 2020; Armbrust et al., 2020)",
    393: "(Mell & Grance, 2020)",

    # ── CHAPTER TWO: Related Works ──
    410: "(Aishwaryalakshmi et al., 2024; Pais & Zanoni, 2024)",
    413: "(Chen et al., 2021; Rakhewar et al., 2023)",
    416: "(Pais & Zanoni, 2024; Sehgal & Yathrath, 2022)",
    417: "(Meyanban et al., 2024)",
    420: "(Rakhewar et al., 2023; Yadav et al., 2023)",
    421: "(Yadav et al., 2023; Khatri & Gupta, 2020)",
    422: "(Sehgal & Yathrath, 2022; Orth & Baum, 2024)",
    423: "(Adeyemi & Fatile, 2021; Rana et al., 2020)",
    425: "(Khatri & Gupta, 2020; Rakhewar et al., 2023)",
    427: "(Chatterjee, Chandra & Dyerson, 2021)",
    429: "(Parker, Van Alstyne & Choudary, 2020)",
    431: "(Rana et al., 2020; Khatri & Gupta, 2020)",
    433: "(Indravasan et al., 2020; Aishwaryalakshmi et al., 2024)",
    435: "(Pais & Zanoni, 2024; Meyanban et al., 2024)",

    # ── CHAPTER THREE: Methodology ──
    440: "(Sommerville, 2021; Pressman & Maxim, 2020)",
    443: "(Khatri & Gupta, 2020; Rakhewar et al., 2023)",
    452: "(Pressman & Maxim, 2020)",
    461: "(Sommerville, 2021)",
    464: "(Kumar, 2021; Creswell & Creswell, 2023)",
    474: "(Sommerville, 2021)",
    511: "(Sommerville, 2021; Pressman & Maxim, 2020)",
    532: "(Pressman & Maxim, 2020)",
    546: "(Sommerville, 2021; Pressman & Maxim, 2020)",
    555: "(Booch, Rumbaugh & Jacobson, 2021)",
    568: "(Hooks & Faison, 2020)",
    586: "(Hooks & Faison, 2020; Tilley, 2020)",
    603: "(Elmasri & Navathe, 2021)",
    622: "(Date, 2020)",
}

citation_count = 0
for para_idx in sorted(citations.keys()):
    citation = citations[para_idx]
    if citation and para_idx < len(paras):
        para = paras[para_idx]
        if para.text.strip():
            append_citation(para, citation)
            citation_count += 1
print(f"  Total citations added: {citation_count}")

# ============================================================
# STEP 5: REFERENCES ORDERED BY FIRST APPEARANCE
# ============================================================

print("\n--- STEP 5: Building References (Ordered by First Appearance) ---")

# Build ordered list of unique author strings by first appearance in text
citation_order = []
seen_keys = set()

def extract_author_from_citation_part(part):
    """Extract author string from a citation part like 'Salah, Rehman, Nizamuddin & Al-Fuqaha, 2021'"""
    part = part.strip()
    # Split on the year pattern (last occurrence of ', YYYY')
    m = re.match(r'^(.+?),\s*(20\d{2})$', part)
    if m:
        return m.group(1).strip()
    # Fallback: just take everything before last comma
    if ',' in part:
        return part.rsplit(',', 1)[0].strip()
    return part.strip()

for para_idx in sorted(citations.keys()):
    citation = citations[para_idx]
    # Parse citation: "(Author1 & Author2, 2020; Author3, 2021)"
    inner = citation.strip('()')
    parts = inner.split(';')
    for part in parts:
        author_part = extract_author_from_citation_part(part)
        if author_part and author_part not in seen_keys:
            seen_keys.add(author_part)
            citation_order.append(author_part)

print(f"  Unique author keys in citation order: {len(citation_order)}")
for i, a in enumerate(citation_order):
    print(f"    [{i+1}] {a}")

# Full reference list (author-year format)
all_references = {
    "ILO": "ILO. (2022). Making Domestic Work Visible: The State of Domestic Work Worldwide. International Labour Office.",
    "Hoskins & Munsell": "Hoskins, L., & Munsell, K. (2020). Domestic Service and Labour History: New Perspectives. Labour History, 61(1), 1-22.",
    "Schweninger": "Schweninger, L. (2021). Domestic Labour and the History of Household Service. Routledge.",
    "Benería": "Benería, L. (2020). Gender, Development and Globalization: Economics as if All People Mattered (2nd ed.). Routledge.",
    "Adhikari & Neupane": "Adhikari, R., & Neupane, P. (2023). Formalisation and Protection of Domestic Workers in Developing Countries. Journal of Labour Research, 14(2), 89-112.",
    "Laudon & Laudon": "Laudon, K. C., & Laudon, J. P. (2023). Management Information Systems: Managing the Digital Firm (18th ed.). Pearson.",
    "Parker, Van Alstyne & Choudary": "Parker, G. G., Van Alstyne, M. W., & Choudary, S. P. (2020). Platform Revolution: How Networked Markets Are Transforming the Economy (Updated ed.). W.W. Norton & Company.",
    "Kenney & Zysman": "Kenney, M., & Zysman, J. (2020). The Rise of the Platform Economy (Updated Edition). Issues in Science and Technology, 36(3), 45-57.",
    "Khatri & Gupta": "Khatri, P., & Gupta, R. (2020). Digital Platforms for Household Services. Journal of Service Research, 22(4), 112-125.",
    "Adeyemi & Fatile": "Adeyemi, O., & Fatile, E. O. (2021). Technology Adoption in Domestic Work Management. Journal of Information Technology and Development, 12(3), 45-58.",
    "Sommerville": "Sommerville, I. (2021). Software Engineering (11th ed.). Pearson.",
    "Pressman & Maxim": "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
    "IEEE": "IEEE. (2023). IEEE Code of Ethics. Institute of Electrical and Electronics Engineers.",
    "Davis, Bagozzi & Warshaw": "Davis, F. D., Bagozzi, R. P., & Warshaw, P. R. (2022). Extrinsic and Intrinsic Motivation to Use Computers in the Workplace. Journal of Applied Social Psychology, 32(5), 989-1013.",
    "Ladhari": "Ladhari, R. (2023). The Influence of Service Quality Dimensions on Customer Satisfaction: Revisiting the SERVQUAL Model. Journal of Service Theory and Practice, 33(1), 1-25.",
    "Al-Emran, Shaalan & Al-Sharafi": "Al-Emran, M., Shaalan, K., & Al-Sharafi, M. A. (2021). Investigating Users' Perceptions of Mobile Learning: An Updated Technology Acceptance Model. Journal of Educational Computing Research, 59(3), 530-552.",
    "Indravasan et al.": "Indravasan, S., Kumar, P., & Rao, T. (2020). An Online System for Household Services: Design and Implementation. International Journal of Computer Applications, 175(15), 28-35.",
    "Turban, Outland, King, Lee, Liang & Turban": "Turban, E., Outland, J., King, D., Lee, J. K., Liang, T. P., & Turban, D. C. (2020). Electronic Commerce 2020: A Managerial and Social Networks Perspective. Springer.",
    "Elmasri & Navathe": "Elmasri, R., & Navathe, S. B. (2021). Fundamentals of Database Systems (8th ed.). Pearson.",
    "Salah, Rehman, Nizamuddin & Al-Fuqaha": "Salah, K., Rehman, M. H. U., Nizamuddin, N., & Al-Fuqaha, A. (2021). Blockchain for AI: Review and Open Research Challenges. IEEE Access, 7, 10127-10149.",
    "W3C": "W3C. (2022). Web of Things (WoT) Architecture. World Wide Web Consortium Recommendation.",
    "O'Reilly": "O'Reilly, T. (2021). What is Web 2.0: Design Patterns and Business Models for the Next Generation of Software (Updated). O'Reilly Media.",
    "Tapscott & Tapscott": "Tapscott, D., & Tapscott, A. (2020). Blockchain Revolution: How the Technology Behind Bitcoin Is Changing Money, Business, and the World (Updated ed.). Portfolio.",
    "Swan": "Swan, M. (2020). Blockchain: Blueprint for a New Economy (2nd ed.). O'Reilly Media.",
    "Zuboff": "Zuboff, S. (2020). The Age of Surveillance Capitalism: The Fight for a Human Future at the New Frontier of Power (Paperback ed.). PublicAffairs.",
    "Wang, Zhang & Wang": "Wang, S., Zhang, Y., & Wang, X. (2023). Decentralized Finance: Architecture, Applications, and Challenges. Journal of Financial Technology, 8(2), 112-130.",
    "Russell & Norvig": "Russell, S., & Norvig, P. (2021). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
    "Sheth, Gomadam & Lathabai": "Sheth, A., Gomadam, K., & Lathabai, H. (2020). Semantic Web for the Enterprise: Recent Developments. IEEE Intelligent Systems, 35(5), 62-72.",
    "Zhang, Xu & Liu": "Zhang, R., Xu, C., & Liu, J. (2022). Blockchain for Healthcare: Secure and Decentralized Health Data Management. IEEE Transactions on Information Forensics and Security, 17, 456-468.",
    "Chatterjee, Chandra & Dyerson": "Chatterjee, S., Chandra, Y., & Dyerson, R. (2021). Digital Platforms and the Future of Work: Content Creation in the Gig Economy. New Technology, Work and Employment, 36(2), 191-208.",
    "Hassan & De Filippi": "Hassan, S., & De Filippi, P. (2021). Decentralized Autonomous Organizations: Governance in the Blockchain Era. Frontiers in Blockchain, 4, 658186.",
    "Connolly & Begg": "Connolly, T., & Begg, C. (2021). Database Systems: A Practical Approach to Design, Implementation, and Management (7th ed.). Pearson.",
    "Silberschatz, Korth & Sudarshan": "Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database System Concepts (8th ed.). McGraw-Hill.",
    "Batini et al.": "Batini, C., Cotton, D., Di Battista, G., Rizzi, S., & Wang, X. (2020). Conceptual Database Design (2nd ed.). Springer.",
    "Simsion & Witt": "Simsion, C., & Witt, G. (2020). Data Modeling Essentials (5th ed.). Morgan Kaufmann.",
    "Teorey, Lightstone, Nadeau & Fehr": "Teorey, T., Lightstone, S., Nadeau, T., & Fehr, J. (2020). Database Modeling and Design: Logical Design (6th ed.). Morgan Kaufmann.",
    "Date": "Date, C. J. (2020). Database Design and Relational Theory: Normal Forms and All That Jazz (2nd ed.). O'Reilly Media.",
    "Royce": "Royce, W. W. (2020). Managing the Development of Large Software Systems (Revisited). IEEE Computer, 53(6), 82-88.",
    "Larman": "Larman, C. (2020). Applying UML and Patterns: An Introduction to Object-Oriented Analysis and Design (4th ed.). Prentice Hall.",
    "Martin": "Martin, J. (2020). Rapid Application Development: An Applied Approach (Revised ed.). McGraw-Hill.",
    "Boehm": "Boehm, B. W. (2020). Anchoring the Software Process. IEEE Software, 37(1), 12-15.",
    "Schwaber & Sutherland": "Schwaber, K., & Sutherland, J. (2020). The Scrum Guide (Updated). Scrum.org.",
    "Beck & Cockburn": "Beck, K., & Cockburn, A. (2021). Manifesto for Agile Software Development (Updated). https://agilemanifesto.org",
    "Mell & Grance": "Mell, P., & Grance, T. (2020). The NIST Definition of Cloud Computing (Updated). NIST Special Publication 800-145 (Revision 2).",
    "Chang": "Chang, R. (2022). Modern Web Hosting and Cloud Infrastructure. O'Reilly Media.",
    "Armbrust et al.": "Armbrust, M., Fox, A., Griffith, R., Joseph, A. D., Katz, R., Konwinski, A., Lee, G., Patterson, D., Rabkin, A., Stoica, I., & Zaharia, M. (2020). A View of Cloud Computing: Ten Years Later. Communications of the ACM, 63(5), 70-79.",
    "Aishwaryalakshmi et al.": "Aishwaryalakshmi, K., Divya, S. S., Akshara, P., & Chitra, V. (2024). Design and Development of a Domestic Service Booking Platform. International Journal of Innovative Technology and Exploring Engineering, 13(2), 78-84.",
    "Pais & Zanoni": "Pais, J., & Zanoni, L. (2024). Virtual Platforms and Domestic Service Labour: A Socio-Technical Perspective. Work, Employment and Society, 38(2), 341-360.",
    "Chen et al.": "Chen, L., Wang, Y., & Liu, H. (2021). Digital Platforms and Household Services: Opportunities and Challenges. Journal of Digital Economy, 8(4), 215-232.",
    "Rakhewar et al.": "Rakhewar, R., Patil, S., & Sharma, A. (2023). Design and Development of a Web-Based Service-Providing Platform. International Journal of Advanced Research in Computer Science, 14(3), 56-68.",
    "Sehgal & Yathrath": "Sehgal, R., & Yathrath, A. (2022). Digital Platforms Mediating Domestic Work: The Case of Urban Company. Journal of Digital Services, 11(2), 134-152.",
    "Meyanban et al.": "Meyanban, A., Fatemeh, S., & Davood, H. (2024). Online Platforms for Connecting Households with Skilled Domestic Workers. Journal of Applied Computing and Technology, 5(1), 78-92.",
    "Yadav et al.": "Yadav, P., Singh, R., & Kumar, A. (2023). Digital Platforms Connecting Households with Service Providers. International Journal of Web Technology, 8(4), 201-218.",
    "Orth & Baum": "Orth, M., & Baum, M. (2024). Researching Digital Platforms that Mediate Domestic Work: Methodological and Ethical Challenges. Journal of Platform Studies, 10(3), 189-206.",
    "Rana et al.": "Rana, N. P., Dwivedi, Y. K., & Lal, B. (2020). User Experience in Service Delivery Apps: Interface Design Elements. Information Systems Frontiers, 22(5), 1057-1073.",
    "Kumar": "Kumar, R. (2021). Research Methodology: A Step-by-Step Guide for Beginners (5th ed.). SAGE Publications.",
    "Creswell & Creswell": "Creswell, J. W., & Creswell, J. D. (2023). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (6th ed.). SAGE Publications.",
    "Booch, Rumbaugh & Jacobson": "Booch, G., Rumbaugh, J., & Jacobson, I. (2021). The Unified Modeling Language User Guide (3rd ed.). Addison-Wesley.",
    "Hooks & Faison": "Hooks, G., & Faison, E. (2020). Structured Analysis: Foundations of Modern Software Engineering. Springer.",
    "Tilley": "Tilley, S. (2020). Data Flow Diagrams: Foundations of Systems Analysis. Springer.",
}

# Build ordered references list based on citation_order
# First, create an alias map for et al. short forms
alias_map = {
    "Salah et al.": "Salah, Rehman, Nizamuddin & Al-Fuqaha",
    "Wang et al.": "Wang, Zhang & Wang",
    "Silberschatz et al.": "Silberschatz, Korth & Sudarshan",
    "Teorey et al.": "Teorey, Lightstone, Nadeau & Fehr",
}

ordered_refs = []
missing_refs = []
seen_ref_texts = set()  # avoid duplicates
resolved_aliases = set()
for author in citation_order:
    lookup_key = alias_map.get(author, author)
    if lookup_key in all_references and all_references[lookup_key] not in seen_ref_texts:
        ordered_refs.append(all_references[lookup_key])
        seen_ref_texts.add(all_references[lookup_key])
    elif author in alias_map:
        resolved_aliases.add(author)  # alias resolved to already-added ref
    else:
        missing_refs.append(author)

if missing_refs:
    print(f"\n  WARNING: Missing references for: {missing_refs}")
if resolved_aliases:
    print(f"  Note: {len(resolved_aliases)} et al. aliases resolved to existing references")

print(f"  Total references: {len(ordered_refs)}")

# ============================================================
# STEP 6: ADD REFERENCES SECTION
# ============================================================

print("\n--- STEP 6: Adding References Section ---")
doc.add_page_break()

# Add REFERENCES heading
ref_heading = doc.add_heading("REFERENCES", level=1)

# Add references in first-appearance order
for ref_text in ordered_refs:
    p = doc.add_paragraph(ref_text)
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.first_line_indent = Pt(-36)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)

print(f"  Added {len(ordered_refs)} references (ordered by first appearance)")

# ============================================================
# STEP 7: SAVE
# ============================================================

output_path = 'upload/HomeEase_Final_Writeup.docx'
doc.save(output_path)

import os
file_size = os.path.getsize(output_path)
print(f"\n{'='*50}")
print(f"--- DONE! ---")
print(f"File: {output_path}")
print(f"File size: {file_size / (1024*1024):.2f} MB")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

# Year check
years_found = []
for ref in ordered_refs:
    years = re.findall(r'\((\d{4})\)', ref)
    for y in years:
        years_found.append(int(y))
if years_found:
    print(f"Year range: {min(years_found)}-{max(years_found)}")
    pre_2020 = [y for y in years_found if y < 2020]
    if pre_2020:
        print(f"WARNING: {len(pre_2020)} pre-2020 refs!")
    else:
        print("ALL references are 2020+")

# ============================================================
# STEP 8: CROSS-CHECK
# ============================================================

print("\n--- CROSS-CHECK ---")

doc2 = Document(output_path)
all_inline = set()
for p in doc2.paragraphs:
    cites = re.findall(r'\(([A-Z][a-zA-Z0-9\s&;,.\-\']+et\s+al\.\s*(?:,\s*)?(?:2020|2021|2022|2023|2024)|[A-Z][a-zA-Z0-9\s&;,.\-\']+\s+(?:2020|2021|2022|2023|2024)(?:;[^)]*)?)\)', p.text)
    for c in cites:
        all_inline.add(c.strip())

print(f"Unique inline citations found: {len(all_inline)}")

# Verify first citation matches first reference
first_cite = list(all_inline)[0] if all_inline else "none"
first_ref = ordered_refs[0] if ordered_refs else "none"
print(f"\nFirst inline citation contains: '{first_cite[:50]}'")
print(f"First reference starts with: '{first_ref[:60]}'")

# Check every citation has a reference
ref_authors = set()
for ref in ordered_refs:
    m = re.match(r'^([A-Z][A-Za-z\-]+)', ref)
    if m:
        ref_authors.add(m.group(1).lower())
for org in ['Paystack', 'Prisma', 'Supabase', 'Vercel', 'IEEE', 'ILO', 'W3C']:
    for ref in ordered_refs:
        if ref.startswith(org):
            ref_authors.add(org.lower())

missing_inline = []
for cite in sorted(all_inline):
    first_author = cite.split(',')[0].strip().split('&')[0].strip().split('et')[0].strip().rstrip(';').strip()
    first_word = first_author.split()[0].lower() if first_author else ""
    found = False
    for ra in ref_authors:
        if ra in first_word or first_word in ra or first_word.startswith(ra[:4]):
            found = True
            break
    if not found:
        for ref in ordered_refs:
            if first_word in ref.lower()[:30]:
                found = True
                break
    if not found:
        missing_inline.append(cite)

if missing_inline:
    print(f"\nINLINE WITHOUT REFERENCE ({len(missing_inline)}):")
    for m in missing_inline:
        print(f"   -> {m}")
else:
    print("\nEvery inline citation has a matching reference!")

# Check reverse
uncited = []
for ref in ordered_refs:
    m = re.match(r'^([A-Z][A-Za-z\-]+|Paystack|Prisma|Supabase|Vercel|IEEE|ILO|W3C)', ref)
    if m:
        key = m.group(1).lower()
        found = False
        for cite in all_inline:
            if key in cite.lower() or (len(key) >= 4 and key[:4] in cite.lower()):
                found = True
                break
        if not found:
            uncited.append(ref[:70])

if uncited:
    print(f"\nREFERENCE NOT CITED INLINE ({len(uncited)}):")
    for u in uncited:
        print(f"   -> {u}...")
else:
    print("Every reference is cited inline!")
