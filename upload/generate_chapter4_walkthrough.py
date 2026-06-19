#!/usr/bin/env python3
"""
Generate HomeEase Chapter 4 - System Implementation and User Guide
A comprehensive step-by-step walkthrough document for the FYP thesis.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==============================================================================
# CONFIGURATION
# ==============================================================================

OUTPUT_PATH = "/home/z/my-project/upload/HomeEase_Chapter_4_Walkthrough.docx"
SCREENSHOTS_DIR = "/home/z/my-project/upload/screenshots"

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
LINE_SPACING = 1.5
MARGIN_INCHES = 1.0

# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing

def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    else:
        run.font.name = FONT_NAME
    # Set East-Asian font
    r_elem = run._element
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r_elem.insert(0, rPr)
    return run

def add_body_paragraph(doc, text, bold=False, italic=False, indent=False, space_after=6):
    """Add a body paragraph with Times New Roman 12pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = BODY_SIZE
    run.bold = bold
    run.italic = italic
    return p

def add_numbered_step(doc, step_number, text, sub_items=None):
    """Add a numbered step with optional sub-items."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.63)

    # Step number and main text
    run_num = p.add_run(f"Step {step_number}: ")
    run_num.font.name = FONT_NAME
    run_num.font.size = BODY_SIZE
    run_num.bold = True

    run_text = p.add_run(text)
    run_text.font.name = FONT_NAME
    run_text.font.size = BODY_SIZE

    # Sub-items as bullet points
    if sub_items:
        for item in sub_items:
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bpf = bp.paragraph_format
            bpf.line_spacing = LINE_SPACING
            bpf.space_after = Pt(2)
            bpf.left_indent = Cm(1.27)

            run_bullet = bp.add_run("• " + item)
            run_bullet.font.name = FONT_NAME
            run_bullet.font.size = BODY_SIZE

def add_mixed_paragraph(doc, parts):
    """Add a paragraph with mixed formatting. parts is a list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
        run.bold = bold
        run.italic = italic
    return p

def add_bullet_point(doc, text, bold_prefix=None, indent_level=0):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(2)
    pf.left_indent = Cm(1.27 + indent_level * 0.63)

    if bold_prefix:
        run_prefix = p.add_run("• " + bold_prefix)
        run_prefix.font.name = FONT_NAME
        run_prefix.font.size = BODY_SIZE
        run_prefix.bold = True
        run_rest = p.add_run(text)
        run_rest.font.name = FONT_NAME
        run_rest.font.size = BODY_SIZE
    else:
        run = p.add_run("• " + text)
        run.font.name = FONT_NAME
        run.font.size = BODY_SIZE
    return p

def add_screenshot(doc, filename, figure_number, caption_text):
    """Add a screenshot image with caption."""
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    
    # Add blank line before image
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(6)
    
    if os.path.exists(filepath):
        # Add image centered, max width 6 inches
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(2)
        p_img.paragraph_format.space_before = Pt(2)
        run_img = p_img.add_run()
        try:
            # Get image dimensions to calculate width
            from PIL import Image as PILImage
            pil_img = PILImage.open(filepath)
            img_width, img_height = pil_img.size
            aspect_ratio = img_height / img_width
            display_width = min(6.0, img_width / 150)  # 6 inches max
            display_height = display_width * aspect_ratio
            if display_height > 4.5:  # max 4.5 inches tall
                display_height = 4.5
                display_width = display_height / aspect_ratio
            run_img.add_picture(filepath, width=Inches(display_width))
        except Exception:
            # Fallback: just add the image at 6 inches
            run_img.add_picture(filepath, width=Inches(6.0))
    else:
        p_warn = doc.add_paragraph()
        p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_warn = p_warn.add_run(f"[Screenshot not found: {filename}]")
        run_warn.font.name = FONT_NAME
        run_warn.font.size = Pt(10)
        run_warn.italic = True
        run_warn.font.color.rgb = RGBColor(192, 0, 0)

    # Add caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    p_cap.paragraph_format.space_before = Pt(2)
    run_cap = p_cap.add_run(f"Figure {figure_number}: {caption_text}")
    run_cap.font.name = FONT_NAME
    run_cap.font.size = Pt(10)
    run_cap.italic = True

def add_table_row(table, cells_data, bold=False, header=False):
    """Add a row to a table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = bold
        if header:
            set_cell_shading(cell, "D9E2F3")
    return row

def add_heading_text(doc, text, level="section"):
    """Add a styled heading."""
    if level == "chapter":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(24)
        pf.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(16)
        run.bold = True
    elif level == "section":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.bold = True
    elif level == "subsection":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.bold = True
    elif level == "subsubsection":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.bold = True
        run.italic = True
    return p


# ==============================================================================
# DOCUMENT GENERATION
# ==============================================================================

def generate_chapter4():
    doc = Document()
    
    # ---- Page Setup ----
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(MARGIN_INCHES)
        section.bottom_margin = Inches(MARGIN_INCHES)
        section.left_margin = Inches(MARGIN_INCHES)
        section.right_margin = Inches(MARGIN_INCHES)
    
    # ---- Set default style ----
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING
    
    # ========================================================================
    # CHAPTER TITLE
    # ========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(48)
    pf.space_after = Pt(6)
    pf.line_spacing = LINE_SPACING
    run = p.add_run("CHAPTER 4")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(24)
    pf.line_spacing = LINE_SPACING
    run = p.add_run("SYSTEM IMPLEMENTATION AND USER GUIDE")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.bold = True

    add_page_break(doc)

    # ========================================================================
    # 4.1 Introduction
    # ========================================================================
    add_heading_text(doc, "4.1 Introduction", "section")

    add_body_paragraph(doc,
        "This chapter presents a comprehensive overview of the HomeEase system implementation "
        "environment and provides a detailed, step-by-step user guide for every process within the "
        "platform. The implementation environment section outlines the hardware and software "
        "requirements that were utilised during the development and deployment of the application. "
        "The user manual section serves as a complete walkthrough covering all user roles: Service "
        "Seekers, Service Providers, and Administrators."
    )

    add_body_paragraph(doc,
        "Each process is described with numbered steps that guide the user from start to finish, "
        "accompanied by screenshots of the actual system interface to ensure clarity and ease of "
        "understanding. The chapter also includes a narrative sample system run that demonstrates a "
        "complete end-to-end transaction on the HomeEase platform, illustrating how all components "
        "of the system work together seamlessly."
    )

    add_body_paragraph(doc,
        "The HomeEase platform is a web-based application accessible at "
        "https://homeease.vercel.app. It connects individuals seeking home maintenance and repair "
        "services with verified artisans and service providers, providing features such as real-time "
        "booking, live job tracking, secure payment processing via Paystack, real-time messaging, "
        "and an admin-managed verification and payout system.",
        indent=True
    )

    # ========================================================================
    # 4.2 Implementation Environment
    # ========================================================================
    add_heading_text(doc, "4.2 Implementation Environment", "section")

    add_body_paragraph(doc,
        "The HomeEase platform was developed and deployed using a carefully selected set of hardware "
        "and software tools. This section documents the development environment, tools, frameworks, "
        "and platforms used to build and deploy the system."
    )

    # ---- 4.2.1 Hardware Requirements ----
    add_heading_text(doc, "4.2.1 Hardware Requirements", "subsection")

    add_body_paragraph(doc,
        "Table 4.1 and Table 4.2 present the hardware specifications used during development and "
        "the deployment server specifications respectively."
    )

    # Development Hardware Table
    p_table_title = doc.add_paragraph()
    p_table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_table_title.paragraph_format.space_after = Pt(4)
    run_tt = p_table_title.add_run("Table 4.1: Development Machine Specifications")
    run_tt.font.name = FONT_NAME
    run_tt.font.size = Pt(10)
    run_tt.bold = True

    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    table1.autofit = True
    # Header row
    hdr = table1.rows[0]
    for i, text in enumerate(["Component", "Specification"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, "D9E2F3")

    dev_hw = [
        ("Processor", "Intel Core i5 / i7 or AMD Ryzen 5 / 7"),
        ("RAM", "16 GB minimum (32 GB recommended)"),
        ("Storage", "256 GB SSD minimum (512 GB recommended)"),
        ("Display", "13-inch or larger (1920 x 1080 minimum)"),
        ("Operating System", "Windows 10/11, macOS 12+, or Ubuntu 20.04+"),
        ("Network", "Stable broadband internet connection"),
    ]
    for comp, spec in dev_hw:
        add_table_row(table1, [comp, spec], bold=False)

    # Set column widths
    for row in table1.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(3.5)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # Deployment Hardware Table
    p_table_title2 = doc.add_paragraph()
    p_table_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_table_title2.paragraph_format.space_after = Pt(4)
    run_tt2 = p_table_title2.add_run("Table 4.2: Deployment Server Specifications (Vercel Serverless)")
    run_tt2.font.name = FONT_NAME
    run_tt2.font.size = Pt(10)
    run_tt2.bold = True

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    table2.autofit = True
    hdr2 = table2.rows[0]
    for i, text in enumerate(["Component", "Specification"]):
        cell = hdr2.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, "D9E2F3")

    dep_hw = [
        ("Platform", "Vercel Serverless Functions"),
        ("Compute", "AWS Lambda-based serverless execution"),
        ("Memory", "Up to 1024 MB per function"),
        ("Execution Timeout", "10 seconds ( Hobby), 60 seconds ( Pro )"),
        ("Database", "Supabase PostgreSQL (managed cloud)"),
        ("CDN", "Vercel Edge Network (global)"),
        ("SSL/TLS", "Automatic HTTPS via Vercel"),
        ("Region", "Global edge deployment"),
    ]
    for comp, spec in dep_hw:
        add_table_row(table2, [comp, spec], bold=False)

    for row in table2.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(3.5)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(12)

    # ---- 4.2.2 Software Requirements ----
    add_heading_text(doc, "4.2.2 Software Requirements", "subsection")

    add_body_paragraph(doc,
        "Table 4.3 presents the comprehensive list of software tools, frameworks, libraries, and "
        "platforms used in the development, testing, and deployment of the HomeEase system."
    )

    p_table_title3 = doc.add_paragraph()
    p_table_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_table_title3.paragraph_format.space_after = Pt(4)
    run_tt3 = p_table_title3.add_run("Table 4.3: Software Requirements")
    run_tt3.font.name = FONT_NAME
    run_tt3.font.size = Pt(10)
    run_tt3.bold = True

    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    table3.autofit = True
    hdr3 = table3.rows[0]
    for i, text in enumerate(["Category", "Tool/Framework", "Version/Purpose"]):
        cell = hdr3.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
        run.bold = True
        set_cell_shading(cell, "D9E2F3")

    sw_reqs = [
        ("Frontend Framework", "Next.js", "v16 (App Router)"),
        ("UI Library", "React", "v19"),
        ("Language", "TypeScript", "v5"),
        ("Styling", "Tailwind CSS", "v4"),
        ("Component Library", "shadcn/ui", "New York style"),
        ("Icon Library", "Lucide React", "Latest"),
        ("Database ORM", "Prisma", "ORM for database management"),
        ("Database", "PostgreSQL", "via Supabase (managed cloud)"),
        ("Real-time Communication", "Socket.IO", "WebSocket-based messaging"),
        ("Authentication", "NextAuth.js", "v4 (credential-based)"),
        ("Payment Processing", "Paystack API", "Secure payments and transfers"),
        ("Runtime", "Bun", "JavaScript/TypeScript runtime"),
        ("Version Control", "Git", "Source code management"),
        ("Code Editor", "VS Code", "Integrated Development Environment"),
        ("API Testing", "Postman", "API endpoint testing"),
        ("Deployment", "Vercel", "Cloud platform (serverless)"),
        ("Package Manager", "Bun", "Fast package installation"),
        ("Animation", "Framer Motion", "UI animations and transitions"),
    ]
    for cat, tool, ver in sw_reqs:
        add_table_row(table3, [cat, tool, ver], bold=False)

    for row in table3.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(1.8)
        row.cells[2].width = Inches(2.2)

    add_page_break(doc)

    # ========================================================================
    # 4.3 System Interface and User Manual
    # ========================================================================
    add_heading_text(doc, "4.3 System Interface and User Manual", "section")

    add_body_paragraph(doc,
        "This section provides a detailed, step-by-step guide for every process within the HomeEase "
        "platform. The user manual is organised by user role and covers all major functionalities "
        "available on the system. Each process is described with clear numbered steps and "
        "accompanied by screenshots of the actual system interface to facilitate understanding."
    )

    # ========================================================================
    # 4.3.1 Getting Started (Common to All Users)
    # ========================================================================
    add_heading_text(doc, "4.3.1 Getting Started (Common to All Users)", "subsection")

    add_body_paragraph(doc,
        "This subsection covers the initial steps that all users of the HomeEase platform must "
        "complete before accessing the system's core functionalities. These include accessing the "
        "website, creating an account, and signing in."
    )

    # ---- 4.3.1.1 Accessing the Website ----
    add_heading_text(doc, "4.3.1.1 Accessing the Website", "subsubsection")

    add_body_paragraph(doc,
        "To access the HomeEase platform, users must have a web browser installed on their device "
        "and an active internet connection. Follow the steps below to access the website:"
    )

    add_numbered_step(doc, 1, "Open a web browser on your device (Google Chrome, Mozilla Firefox, Microsoft Edge, or Safari are recommended).")
    add_numbered_step(doc, 2, 'In the address bar, type the following URL: https://homeease.vercel.app and press Enter.')
    add_numbered_step(doc, 3, "The HomeEase homepage loads, displaying the platform's hero section, featured services, how-it-works section, and navigation links for signing up and signing in.")

    add_screenshot(doc, "01_landing_page.png", "4.1", "HomeEase Landing Page")

    add_body_paragraph(doc,
        "The landing page (Figure 4.1) presents an overview of the platform's capabilities, "
        "including the main value proposition, available service categories, and call-to-action "
        "buttons for user registration. Users can scroll through the page to learn more about "
        "the platform before creating an account."
    )

    # ---- 4.3.1.2 Creating an Account ----
    add_heading_text(doc, "4.3.1.2 Creating an Account", "subsubsection")

    add_body_paragraph(doc,
        "HomeEase requires all users to create an account before accessing platform features. "
        "The registration process begins with role selection, as the platform caters to two "
        "distinct user types: Service Seekers (clients) and Service Providers (artisans)."
    )

    add_numbered_step(doc, 1, 'Click the "Get Started" button prominently displayed on the homepage navigation bar or hero section.')
    add_numbered_step(doc, 2, 'The Role Selection screen appears, presenting two options: "Service Seeker" and "Service Provider". Each option is displayed as a card with a brief description of the role.')

    add_screenshot(doc, "04_role_selection.png", "4.2", "Role Selection – Service Seeker or Service Provider")

    add_body_paragraph(doc,
        "Figure 4.2 shows the role selection interface where new users choose their account type. "
        "This distinction ensures that each user is presented with the appropriate registration "
        "form and dashboard features tailored to their role."
    )

    # -- Signing Up as a Service Seeker --
    add_body_paragraph(doc, "Signing Up as a Service Seeker:", bold=True)

    add_numbered_step(doc, 1, 'Click the "Service Seeker" card on the role selection screen.')
    add_numbered_step(doc, 2, "The Service Seeker registration form appears. Fill in the required fields:", sub_items=[
        "Full Name – Enter your first and last name",
        "Email Address – Provide a valid email address for account verification",
        "Phone Number – Enter your active phone number",
        "Password – Create a strong password (minimum 8 characters)",
        "Confirm Password – Re-enter the password to confirm",
    ])
    add_numbered_step(doc, 3, 'Click the "Sign Up" button at the bottom of the form.')
    add_numbered_step(doc, 4, "If all fields are correctly filled, your account is created and you are redirected to the login page. You can now sign in with your credentials.")

    add_screenshot(doc, "05_seeker_signup.png", "4.3", "Service Seeker Registration Form")

    add_body_paragraph(doc,
        "Figure 4.3 shows the Service Seeker registration form. The form collects basic personal "
        "information required to create the user account. All fields marked with an asterisk (*) are "
        "mandatory and must be completed before submission."
    )

    # -- Signing Up as a Service Provider --
    add_body_paragraph(doc, "Signing Up as a Service Provider:", bold=True)

    add_numbered_step(doc, 1, 'Click the "Service Provider" card on the role selection screen.')
    add_numbered_step(doc, 2, "The Service Provider registration form appears. Fill in the required fields:", sub_items=[
        "Full Name – Enter your first and last name",
        "Email Address – Provide a valid email address",
        "Phone Number – Enter your active phone number",
        "Skills – Type each skill (e.g., Plumbing, Electrical, Cleaning) and press Enter to add it as a tag",
        "Bio – Write a brief professional description of your services and experience",
        "Hourly Rate – Enter your service charge rate (in Naira)",
        "Location – Enter your city or area of operation",
        "Bank Name – Select your bank from the dropdown",
        "Account Number – Enter your bank account number",
        "Account Name – Enter the account holder's name as registered with the bank",
        "Password – Create a strong password",
        "Confirm Password – Re-enter the password to confirm",
    ])
    add_numbered_step(doc, 3, 'Click the "Sign Up" button at the bottom of the form.')
    add_numbered_step(doc, 4, 'Your account is created with a status of "Pending Review". An administrator must verify your account before you can access the full provider dashboard.')
    add_numbered_step(doc, 5, "You will receive a notification once your account has been verified and approved by the administrator.")

    add_screenshot(doc, "07_provider_signup.png", "4.4", "Service Provider Registration Form")

    add_body_paragraph(doc,
        "Figure 4.4 shows the Service Provider registration form, which collects additional "
        "professional details such as skills, bio, hourly rate, location, and bank account "
        "information. These details are necessary for the platform to match providers with "
        "service requests and to process payments."
    )

    add_page_break(doc)

    # ---- 4.3.1.3 Signing In ----
    add_heading_text(doc, "4.3.1.3 Signing In", "subsubsection")

    add_body_paragraph(doc,
        "After creating an account, users must sign in to access their personalised dashboard "
        "and platform features. The sign-in process is straightforward and common to all user roles."
    )

    add_numbered_step(doc, 1, 'Click the "Sign In" link on the homepage navigation bar.')
    add_numbered_step(doc, 2, "The login form appears. Enter your credentials:", sub_items=[
        "Email Address, Username, or Phone Number",
        "Password",
    ])
    add_numbered_step(doc, 3, 'Click the "Sign In" button to authenticate.')
    add_numbered_step(doc, 4, "Upon successful authentication, you are redirected to your role-specific dashboard.")

    add_screenshot(doc, "06_login.png", "4.5", "HomeEase Login Form")

    add_body_paragraph(doc,
        "Figure 4.5 shows the login form where existing users enter their credentials to "
        "access the platform. The system supports login with email address, username, or phone "
        "number along with the account password."
    )

    add_page_break(doc)

    # ========================================================================
    # 4.3.2 Service Seeker User Manual
    # ========================================================================
    add_heading_text(doc, "4.3.2 Service Seeker User Manual", "subsection")

    add_body_paragraph(doc,
        "This subsection provides a comprehensive guide for Service Seekers (clients) on how to "
        "use the HomeEase platform. It covers all features available to registered seekers, including "
        "dashboard navigation, finding and booking artisans, managing requests, making payments, "
        "rating providers, messaging, notifications, and profile management."
    )

    # ---- 4.3.2.1 Viewing the Dashboard ----
    add_heading_text(doc, "4.3.2.1 Viewing the Dashboard", "subsubsection")

    add_body_paragraph(doc,
        "After signing in successfully, the Service Seeker is presented with a personalised "
        "dashboard that serves as the central hub for all platform interactions. The dashboard "
        "features a sidebar navigation menu with the following sections:"
    )

    add_bullet_point(doc, " Overview – Displays summary statistics including total service requests, active jobs, pending payments, and average ratings.")
    add_bullet_point(doc, " Find Artisans – Navigate to the artisan search and booking page.")
    add_bullet_point(doc, " My Requests – View and manage all service requests with their current statuses.")
    add_bullet_point(doc, " Payments – Access payment history and initiate payments for completed services.")
    add_bullet_point(doc, " Messages – Access real-time chat conversations with assigned service providers.")
    add_bullet_point(doc, " Profile – View and manage personal account details and preferences.")
    add_bullet_point(doc, " Support – Access the support chat for assistance from the platform administrator.")

    add_screenshot(doc, "08_client_overview.png", "4.6", "Service Seeker Dashboard Overview")

    add_body_paragraph(doc,
        "Figure 4.6 shows the Service Seeker dashboard overview with key metrics and quick-access "
        "navigation. The dashboard provides an at-a-glance summary of the user's activity on the "
        "platform, including the number of active requests, pending payments, and recent notifications."
    )

    # ---- 4.3.2.2 Finding and Booking a Service Provider ----
    add_heading_text(doc, "4.3.2.2 Finding and Booking a Service Provider", "subsubsection")

    add_body_paragraph(doc,
        "One of the core features of HomeEase is the ability for Service Seekers to search for, "
        "browse, and book verified service providers. The following steps describe this process:"
    )

    add_numbered_step(doc, 1, 'On the dashboard, click "Find Artisans" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "The artisan search page opens, displaying a grid of registered and verified service providers.")
    add_numbered_step(doc, 3, 'Use the search bar at the top to search by service type (e.g., "Plumber", "Electrician", "Cook", "Cleaner").')
    add_numbered_step(doc, 4, "Browse the list of matching artisans. Each artisan card displays:", sub_items=[
        "Provider name and profile photo",
        "Average rating (out of 5 stars)",
        "List of skills and services offered",
        "Hourly rate",
        "Location",
    ])
    add_numbered_step(doc, 5, 'Click the "Book Now" button on the desired artisan\'s card.')
    add_numbered_step(doc, 6, "The booking form appears. Fill in the following details:", sub_items=[
        "Location – Enter the address where the service is needed",
        "Preferred Date – Select the date for the service appointment",
        "Preferred Time – Select the preferred time slot",
        "Description – Optionally describe the specific service needed",
    ])
    add_numbered_step(doc, 7, 'Click "Confirm Booking" to submit the request.')
    add_numbered_step(doc, 8, "A confirmation message appears, and the service request is sent to the selected provider. The request status is set to PENDING.")

    add_body_paragraph(doc,
        "The provider will receive the booking request on their dashboard and can choose to accept "
        "or decline it. The seeker is notified of the provider's decision via the platform's "
        "notification system."
    )

    add_screenshot(doc, "03_services_grid.png", "4.7", "Find Artisans – Service Provider Search Grid")

    add_body_paragraph(doc,
        "Figure 4.7 shows the artisan search interface where Service Seekers can browse verified "
        "providers and their profiles. The grid layout makes it easy to compare providers based on "
        "ratings, skills, and pricing."
    )

    # ---- 4.3.2.3 Viewing My Requests ----
    add_heading_text(doc, "4.3.2.3 Viewing My Requests", "subsubsection")

    add_body_paragraph(doc,
        "The My Requests section allows Service Seekers to track the status of all service "
        "requests they have submitted. Follow these steps to view and manage requests:"
    )

    add_numbered_step(doc, 1, 'Click "My Requests" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "A list of all your service requests is displayed. Each request card shows:", sub_items=[
        "Service type requested",
        "Assigned provider name (if applicable)",
        "Current status (PENDING, ACCEPTED, IN_PROGRESS, COMPLETED, AWAITING_PAYMENT, or CANCELLED)",
        "Scheduled date and time",
        "Payment status",
    ])
    add_numbered_step(doc, 3, "Click on a specific request card to view full details, initiate a chat with the provider, or make a payment.")
    add_numbered_step(doc, 4, "Use the status filter tabs at the top to filter requests by their current status.")

    # ---- 4.3.2.4 Making Payment ----
    add_heading_text(doc, "4.3.2.4 Making Payment", "subsubsection")

    add_body_paragraph(doc,
        "HomeEase integrates with Paystack to provide a secure and seamless payment experience. "
        "Service Seekers can make payments for completed services directly through the platform. "
        "The payment process is as follows:"
    )

    add_numbered_step(doc, 1, 'On "My Requests", locate a service request that is marked as COMPLETED or AWAITING_PAYMENT.')
    add_numbered_step(doc, 2, "Click on the request card to open the request details.")
    add_numbered_step(doc, 3, 'Click the "Pay Now" button on the request details page.')
    add_numbered_step(doc, 4, "Enter the amount to pay in the payment form.")
    add_numbered_step(doc, 5, 'Click "Proceed to Payment" to initiate the transaction.')
    add_numbered_step(doc, 6, "You are redirected to the Paystack secure payment page.")
    add_numbered_step(doc, 7, "Enter your card details (card number, expiry date, CVV) or select an alternative payment method (bank transfer, USSD).")
    add_numbered_step(doc, 8, "Complete the payment on the Paystack page.")
    add_numbered_step(doc, 9, "After successful payment, you are redirected back to HomeEase. The payment status on the request is automatically updated to PAID.")
    add_numbered_step(doc, 10, "You may now rate the service provider and leave a review.")

    add_body_paragraph(doc,
        "All payments are processed through Paystack's secure infrastructure, ensuring that "
        "sensitive financial information is protected. The platform retains a 5% commission on "
        "each transaction, and the remaining amount is credited to the service provider's wallet."
    )

    # ---- 4.3.2.5 Rating a Service Provider ----
    add_heading_text(doc, "4.3.2.5 Rating a Service Provider", "subsubsection")

    add_body_paragraph(doc,
        "After a completed service and successful payment, Service Seekers are encouraged to rate "
        "the service provider. Ratings help maintain service quality on the platform and assist "
        "other users in making informed decisions."
    )

    add_numbered_step(doc, 1, "After making payment for a completed service, a rating prompt appears on the request details page.")
    add_numbered_step(doc, 2, "Click on the star rating to assign a score from 1 to 5 stars (1 being poor, 5 being excellent).")
    add_numbered_step(doc, 3, "Optionally, write a review comment describing your experience with the provider.")
    add_numbered_step(doc, 4, 'Click "Submit Review" to save your rating and review.')
    add_numbered_step(doc, 5, "The provider's overall rating is updated based on your feedback.")

    add_body_paragraph(doc,
        "Ratings are visible on the provider's profile and in search results, helping other "
        "seekers choose quality service providers. Providers with consistently low ratings may "
        "be flagged for review by the platform administrator."
    )

    # ---- 4.3.2.6 Messaging a Service Provider ----
    add_heading_text(doc, "4.3.2.6 Messaging a Service Provider", "subsubsection")

    add_body_paragraph(doc,
        "HomeEase provides a real-time messaging system powered by Socket.IO, enabling direct "
        "communication between Service Seekers and their assigned Service Providers. To use the "
        "messaging feature:"
    )

    add_numbered_step(doc, 1, 'Click "Messages" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "A list of your conversations with service providers appears on the left panel.")
    add_numbered_step(doc, 3, "Click on a conversation to open the chat window on the right.")
    add_numbered_step(doc, 4, "Type your message in the text input field at the bottom of the chat window.")
    add_numbered_step(doc, 5, "Click the send button (arrow icon) or press Enter to send the message.")
    add_numbered_step(doc, 6, "Messages are delivered in real-time via WebSocket connection.")
    add_numbered_step(doc, 7, "The provider's replies appear in the chat window instantly.")

    add_body_paragraph(doc,
        "The real-time messaging feature ensures that seekers and providers can communicate "
        "efficiently about service details, scheduling, and any other concerns without leaving the "
        "platform. All messages are associated with a specific service request for context."
    )

    # ---- 4.3.2.7 Managing Notifications ----
    add_heading_text(doc, "4.3.2.7 Managing Notifications", "subsubsection")

    add_body_paragraph(doc,
        "HomeEase provides a notification system to keep users informed about important events "
        "and updates on the platform. Notifications include new service request responses, status "
        "changes, payment confirmations, and new messages."
    )

    add_numbered_step(doc, 1, "Click the bell icon located in the top navigation bar of the dashboard.")
    add_numbered_step(doc, 2, "A dropdown panel appears displaying your recent notifications.")
    add_numbered_step(doc, 3, "Unread notifications are indicated by a count badge on the bell icon.")
    add_numbered_step(doc, 4, "Click on a notification to be redirected to the relevant page (e.g., a new message notification takes you to the Messages section).")
    add_numbered_step(doc, 5, "Notifications are automatically marked as read once clicked.")

    add_body_paragraph(doc,
        "The notification system ensures that users never miss important updates about their "
        "service requests, payments, or messages. Notifications persist in the system and can be "
        "reviewed at any time."
    )

    # ---- 4.3.2.8 Managing Profile ----
    add_heading_text(doc, "4.3.2.8 Managing Profile", "subsubsection")

    add_body_paragraph(doc,
        "Service Seekers can view and update their personal information through the Profile "
        "section of the dashboard."
    )

    add_numbered_step(doc, 1, 'Click "Profile" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "View your current profile details, including name, email, phone number, and role.")
    add_numbered_step(doc, 3, 'Click "Edit Profile" to enter edit mode.')
    add_numbered_step(doc, 4, "Modify any fields as needed (name, email, phone number, password).")
    add_numbered_step(doc, 5, 'Click "Save Changes" to update your profile.')

    add_body_paragraph(doc,
        "Profile information is used for communication and service request purposes. It is "
        "recommended to keep profile details up to date to ensure seamless interactions with "
        "service providers."
    )

    # ---- 4.3.2.9 Using Support Chat ----
    add_heading_text(doc, "4.3.2.9 Using Support Chat", "subsubsection")

    add_body_paragraph(doc,
        "HomeEase provides a built-in support chat feature that allows Service Seekers to "
        "communicate directly with the platform administrator for assistance with any issues "
        "or questions."
    )

    add_numbered_step(doc, 1, 'From the dashboard, click "Support" in the sidebar navigation.')
    add_numbered_step(doc, 2, "The support chat interface opens, displaying any previous conversations.")
    add_numbered_step(doc, 3, "Type your question or describe the issue you are experiencing in the message field.")
    add_numbered_step(doc, 4, "Click the send button to submit your message.")
    add_numbered_step(doc, 5, "The administrator responds to your message in real-time.")
    add_numbered_step(doc, 6, "Continue the conversation until your issue is resolved.")

    add_body_paragraph(doc,
        "The support chat is accessible at all times and provides a convenient channel for "
        "resolving issues related to bookings, payments, account management, or any other "
        "platform-related concerns."
    )

    add_page_break(doc)

    # ========================================================================
    # 4.3.3 Service Provider User Manual
    # ========================================================================
    add_heading_text(doc, "4.3.3 Service Provider User Manual", "subsection")

    add_body_paragraph(doc,
        "This subsection provides a detailed guide for Service Providers (artisans) on how to use "
        "the HomeEase platform. It covers the verification process, dashboard features, job management, "
        "check-in/check-out, transactions, earnings, messaging, and profile management."
    )

    # ---- 4.3.3.1 Awaiting Verification ----
    add_heading_text(doc, "4.3.3.1 Awaiting Verification", "subsubsection")

    add_body_paragraph(doc,
        "Newly registered Service Providers must undergo a verification process before they can "
        "fully access the platform. This process ensures that only legitimate and qualified "
        "providers offer services on HomeEase."
    )

    add_numbered_step(doc, 1, "After completing registration, sign in with your email and password.")
    add_numbered_step(doc, 2, 'A notification banner appears on your dashboard stating: "Your account is pending verification."')
    add_numbered_step(doc, 3, "During the pending period, your dashboard shows limited functionality – you can view your profile but cannot access job offers or messaging.")
    add_numbered_step(doc, 4, "Wait for the platform administrator to review and verify your account. This typically occurs within 24 hours of registration.")
    add_numbered_step(doc, 5, "Once verified, the page automatically updates to display the full provider dashboard with all features enabled.")
    add_numbered_step(doc, 6, "A confirmation notification appears, and you can now receive and manage job offers.")

    add_body_paragraph(doc,
        "If your application is rejected, you will receive a notification with the reason for "
        "rejection and may contact support for further assistance or re-apply after addressing "
        "the concerns."
    )

    # ---- 4.3.3.2 Viewing the Provider Dashboard ----
    add_heading_text(doc, "4.3.3.2 Viewing the Provider Dashboard", "subsubsection")

    add_body_paragraph(doc,
        "After verification, the Service Provider is presented with a comprehensive dashboard "
        "designed for managing job offers, active jobs, transactions, earnings, and communications. "
        "The sidebar navigation contains the following sections:"
    )

    add_bullet_point(doc, " Job Offers – View and manage incoming service requests from seekers.")
    add_bullet_point(doc, " My Jobs – Track currently active and accepted jobs with live timer functionality.")
    add_bullet_point(doc, " Transactions – View all payment transactions with detailed breakdowns.")
    add_bullet_point(doc, " Earnings – Monitor total earnings, available balance, pending amounts, and withdrawal history.")
    add_bullet_point(doc, " Messages – Access real-time chat conversations with clients.")
    add_bullet_point(doc, " Profile – Manage account details, skills, bio, hourly rate, location, and bank information.")
    add_bullet_point(doc, " Support – Access the support chat for platform-related assistance.")

    add_screenshot(doc, "02_dashboard.png", "4.8", "Service Provider Dashboard")

    add_body_paragraph(doc,
        "Figure 4.8 shows the Service Provider dashboard with its comprehensive set of tools "
        "for managing jobs, tracking earnings, and communicating with clients."
    )

    # ---- 4.3.3.3 Managing Job Offers ----
    add_heading_text(doc, "4.3.3.3 Managing Job Offers", "subsubsection")

    add_body_paragraph(doc,
        "When a Service Seeker books a provider, the request appears as a job offer on the "
        "provider's dashboard. Providers can review and either accept or decline incoming offers."
    )

    add_numbered_step(doc, 1, 'Click "Job Offers" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "View the list of incoming service requests. Each offer displays:", sub_items=[
        "Service type requested",
        "Client name and location",
        "Preferred date and time",
        "Description of the service needed",
    ])
    add_numbered_step(doc, 3, "Review each offer carefully to determine if you can fulfil the request.")
    add_numbered_step(doc, 4, 'Click "Accept" to take the job, or "Decline" to reject the offer.')
    add_numbered_step(doc, 5, "After accepting, the job moves to the \"My Jobs\" section and the status changes to ACCEPTED.")
    add_numbered_step(doc, 6, "A notification is automatically sent to the client confirming your acceptance.")
    add_numbered_step(doc, 7, "The client can now view your acceptance in their request details and begin communicating with you via the messaging feature.")

    # ---- 4.3.3.4 Checking In to a Job ----
    add_heading_text(doc, "4.3.3.4 Checking In to a Job", "subsubsection")

    add_body_paragraph(doc,
        "The check-in feature enables Service Providers to start a live timer when they arrive "
        "at a client's location to begin work. This provides transparency and accurate time "
        "tracking for service delivery."
    )

    add_numbered_step(doc, 1, 'After arriving at the client\'s location, click "My Jobs" in the sidebar.')
    add_numbered_step(doc, 2, "Locate the accepted job for which you have arrived.")
    add_numbered_step(doc, 3, 'Click the "Check In" button on the job card.')
    add_numbered_step(doc, 4, "A live timer starts, displaying the elapsed time in HH:MM:SS format.")
    add_numbered_step(doc, 5, 'A "LIVE" indicator confirms that the timer is actively running.')
    add_numbered_step(doc, 6, "The client can see the live timer on their end as well, providing real-time visibility of the service progress.")
    add_numbered_step(doc, 7, "Begin working on the service. The timer continues running in the background.")

    add_body_paragraph(doc,
        "The live timer feature enhances trust between seekers and providers by providing "
        "transparent time tracking. Both parties can monitor the elapsed time throughout the "
        "service delivery."
    )

    # ---- 4.3.3.5 Checking Out After Job Completion ----
    add_heading_text(doc, "4.3.3.5 Checking Out After Job Completion", "subsubsection")

    add_body_paragraph(doc,
        "When the service has been completed, the provider must check out to stop the timer "
        "and mark the job as complete. This triggers the payment process."
    )

    add_numbered_step(doc, 1, "When the work is finished, click the \"Check Out\" button on the active job in the My Jobs section.")
    add_numbered_step(doc, 2, "The live timer stops and records the total duration of the service.")
    add_numbered_step(doc, 3, "The job status automatically changes to \"COMPLETED\".")
    add_numbered_step(doc, 4, "A notification is sent to the client informing them that the service has been completed.")
    add_numbered_step(doc, 5, "The client reviews the completed work and proceeds to make payment through the platform.")
    add_numbered_step(doc, 6, "Wait for the client to confirm payment. You will receive a notification once the payment is successful.")

    add_body_paragraph(doc,
        "The check-out process creates a complete record of the service delivery, including "
        "the total time spent, which is used for payment calculations and service history."
    )

    # ---- 4.3.3.6 Viewing Transactions ----
    add_heading_text(doc, "4.3.3.6 Viewing Transactions", "subsubsection")

    add_body_paragraph(doc,
        "The Transactions section provides Service Providers with a detailed record of all "
        "payment transactions related to their services."
    )

    add_numbered_step(doc, 1, 'Click "Transactions" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "A list of all your transactions appears. Each transaction record displays:", sub_items=[
        "Client name",
        "Service type",
        "Client paid (gross amount before commission)",
        "Platform fee (5% commission deducted)",
        "Your payout (net amount after commission)",
        "Status (Pending, Escrow, Paid, or Refunded)",
        "Transaction date",
    ])
    add_numbered_step(doc, 3, 'Use the filter tabs at the top (All, Paid, Escrow, Refunded) to filter transactions by status.')
    add_numbered_step(doc, 4, "Summary cards at the top of the page display totals for gross earnings, platform fees, and net payouts.")

    add_body_paragraph(doc,
        "The transaction history provides full transparency into all financial activities "
        "on the platform, allowing providers to track their income and verify deductions."
    )

    # ---- 4.3.3.7 Managing Earnings and Wallet ----
    add_heading_text(doc, "4.3.3.7 Managing Earnings and Wallet", "subsubsection")

    add_body_paragraph(doc,
        "The Earnings section allows Service Providers to track their cumulative earnings, "
        "available wallet balance, and initiate withdrawal requests."
    )

    add_numbered_step(doc, 1, 'Click "Earnings" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "View your earnings summary dashboard, which displays:", sub_items=[
        "Total Earnings – Cumulative income from all completed services",
        "Available Balance – Funds available for withdrawal",
        "Pending Amount – Income from recently completed services awaiting clearance",
        "Withdrawn Amount – Total funds already withdrawn to your bank account",
    ])
    add_numbered_step(doc, 3, "To withdraw funds from your wallet:", sub_items=[
        'Click the "Withdraw" button',
        "Enter the withdrawal amount (must not exceed available balance)",
        "Confirm the withdrawal request",
    ])
    add_numbered_step(doc, 4, "The withdrawal request is submitted to the platform administrator for processing.")
    add_numbered_step(doc, 5, "The administrator processes the payout and transfers the funds to your registered bank account via Paystack Transfer.")
    add_numbered_step(doc, 6, "You receive a notification when the payout has been completed successfully.")
    add_numbered_step(doc, 7, "Your earnings summary and available balance are updated accordingly.")

    add_body_paragraph(doc,
        "Withdrawal processing times may vary depending on bank transfer schedules. The "
        "platform administrator reviews all withdrawal requests to prevent fraud and ensure "
        "correct fund transfers."
    )

    # ---- 4.3.3.8 Messaging Clients ----
    add_heading_text(doc, "4.3.3.8 Messaging Clients", "subsubsection")

    add_body_paragraph(doc,
        "Service Providers can communicate with their clients in real-time using the built-in "
        "messaging system. This feature is essential for discussing service details, scheduling, "
        "and any clarifications needed during the service delivery process."
    )

    add_numbered_step(doc, 1, 'Click "Messages" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "A list of conversations with clients appears on the left panel.")
    add_numbered_step(doc, 3, "Click on a conversation to open the chat window.")
    add_numbered_step(doc, 4, "Type your message in the text input field at the bottom of the chat.")
    add_numbered_step(doc, 5, "Click the send button or press Enter to deliver the message.")
    add_numbered_step(doc, 6, "Messages are transmitted in real-time via Socket.IO WebSocket connection.")
    add_numbered_step(doc, 7, "Client replies appear in the chat window instantly, enabling efficient two-way communication.")

    # ---- 4.3.3.9 Managing Profile ----
    add_heading_text(doc, "4.3.3.9 Managing Profile", "subsubsection")

    add_body_paragraph(doc,
        "Service Providers can update their professional profile information, including skills, "
        "bio, rates, and banking details, through the Profile section."
    )

    add_numbered_step(doc, 1, 'Click "Profile" in the sidebar navigation menu.')
    add_numbered_step(doc, 2, "View your current profile details including name, email, phone, skills, bio, hourly rate, location, and bank account information.")
    add_numbered_step(doc, 3, 'Click "Edit Profile" to enter edit mode.')
    add_numbered_step(doc, 4, "Update the following fields as needed:", sub_items=[
        "Skills – Add new skills by typing and pressing Enter; remove existing skills by clicking the X button on the tag",
        "Bio – Update your professional description",
        "Hourly Rate – Adjust your service charge rate",
        "Location – Update your area of operation",
        "Bank Account Details – Update bank name, account number, or account name",
    ])
    add_numbered_step(doc, 5, 'Click "Save Changes" to update your profile.')
    add_numbered_step(doc, 6, "Your updated profile information is immediately reflected on your public profile visible to Service Seekers.")

    add_body_paragraph(doc,
        "Keeping your profile information accurate and up to date helps attract more clients "
        "and ensures smooth payment processing. It is particularly important to maintain correct "
        "bank account details to avoid payment issues."
    )

    add_page_break(doc)

    # ========================================================================
    # 4.3.4 Administrator User Manual
    # ========================================================================
    add_heading_text(doc, "4.3.4 Administrator User Manual", "subsection")

    add_body_paragraph(doc,
        "The Administrator (admin) has the highest level of access on the HomeEase platform "
        "and is responsible for managing the overall operation of the system. This subsection "
        "covers all administrative functions including provider verification, user management, "
        "payout processing, platform statistics, and support handling."
    )

    # ---- 4.3.4.1 Accessing the Admin Dashboard ----
    add_heading_text(doc, "4.3.4.1 Accessing the Admin Dashboard", "subsubsection")

    add_body_paragraph(doc,
        "The admin dashboard provides a comprehensive overview of the entire platform's operations "
        "and offers tools for managing users, requests, payments, and support."
    )

    add_numbered_step(doc, 1, "Sign in with your administrator credentials (email and password).")
    add_numbered_step(doc, 2, "The admin dashboard loads, displaying:", sub_items=[
        "Platform statistics overview (total users, providers, requests, revenue)",
        "Sidebar navigation with management options",
    ])
    add_numbered_step(doc, 3, "The sidebar contains the following management sections:", sub_items=[
        "Overview – Key platform metrics and statistics",
        "Users – Manage all registered users (clients and providers)",
        "Requests – View and manage all service requests",
        "Payouts – Process withdrawal requests from providers",
        "Support – Handle support chat conversations",
        "Settings – Configure platform settings",
    ])

    # ---- 4.3.4.2 Verifying Service Providers ----
    add_heading_text(doc, "4.3.4.2 Verifying Service Providers", "subsubsection")

    add_body_paragraph(doc,
        "One of the most critical administrative functions is the verification of new Service "
        "Provider accounts. This process ensures that only legitimate and qualified artisans "
        "can offer services on the platform."
    )

    add_numbered_step(doc, 1, 'Click "Users" in the admin sidebar navigation.')
    add_numbered_step(doc, 2, 'Use the "Pending" filter tab to display only unverified provider accounts.')
    add_numbered_step(doc, 3, "Click on a pending provider to view their complete registration details.")
    add_numbered_step(doc, 4, "Review the provider's submitted information:", sub_items=[
        "Name and contact information (email, phone)",
        "Skills and professional bio",
        "Bank account details",
    ])
    add_numbered_step(doc, 5, 'Click "Verify" to approve the provider, OR click "Reject" to decline with an optional reason.')
    add_numbered_step(doc, 6, "The provider receives an automatic notification of the verification decision.")
    add_numbered_step(doc, 7, "Verified providers gain immediate access to the full provider dashboard and can begin receiving job offers.")

    add_screenshot(doc, "07_final_verify.png", "4.9", "Admin Dashboard – Provider Verification")

    add_body_paragraph(doc,
        "Figure 4.9 shows the admin interface for reviewing and verifying service provider "
        "applications. The admin can assess each provider's qualifications and make an informed "
        "decision before granting platform access."
    )

    # ---- 4.3.4.3 Managing Users ----
    add_heading_text(doc, "4.3.4.3 Managing Users", "subsubsection")

    add_body_paragraph(doc,
        "The admin can view, search, and manage all registered users on the platform."
    )

    add_numbered_step(doc, 1, 'Click "Users" in the admin sidebar navigation.')
    add_numbered_step(doc, 2, "View the complete list of registered users with their basic information.")
    add_numbered_step(doc, 3, "Use the search bar to find specific users by name, email address, or phone number.")
    add_numbered_step(doc, 4, "Apply filters to view users by role (Client, Provider, or Pending Verification).")
    add_numbered_step(doc, 5, "Click on a user to view their full profile, including activity history, service requests, and transaction records.")
    add_numbered_step(doc, 6, "The admin can take appropriate actions such as suspending accounts or contacting users regarding platform issues.")

    # ---- 4.3.4.4 Processing Payouts ----
    add_heading_text(doc, "4.3.4.4 Processing Payouts", "subsubsection")

    add_body_paragraph(doc,
        "When Service Providers request withdrawals from their earnings wallet, the admin must "
        "process these payout requests to transfer funds to the providers' bank accounts."
    )

    add_numbered_step(doc, 1, 'Click "Payouts" in the admin sidebar navigation.')
    add_numbered_step(doc, 2, "View the list of pending payout requests. Each request shows:", sub_items=[
        "Provider name and contact details",
        "Service amount (total charged to client)",
        "Platform commission (5% deduction)",
        "Net payout amount (amount to be transferred)",
        "Provider's bank account details (bank name, account number, account name)",
        "Request date",
    ])
    add_numbered_step(doc, 3, 'Click "Initiate Payout" on a pending request.')
    add_numbered_step(doc, 4, "Review and confirm the payout details in the confirmation dialog.")
    add_numbered_step(doc, 5, "The system processes the payout via Paystack Transfer API.")
    add_numbered_step(doc, 6, "The provider receives the funds directly in their registered bank account.")
    add_numbered_step(doc, 7, "The transaction status updates to \"Paid\" and the provider's wallet balance is adjusted accordingly.")
    add_numbered_step(doc, 8, "The provider receives a notification confirming the successful payout.")

    add_body_paragraph(doc,
        "Payout processing is a critical financial operation. The admin must verify all payout "
        "details before initiating transfers to prevent errors and ensure funds reach the correct "
        "bank accounts."
    )

    # ---- 4.3.4.5 Viewing Platform Statistics ----
    add_heading_text(doc, "4.3.4.5 Viewing Platform Statistics", "subsubsection")

    add_body_paragraph(doc,
        "The admin dashboard provides real-time statistics about the platform's performance and "
        "growth. These metrics help the administrator monitor the health of the platform and make "
        "informed decisions."
    )

    add_numbered_step(doc, 1, 'Click "Overview" or "Dashboard" in the admin sidebar.')
    add_numbered_step(doc, 2, "View the key performance metrics displayed on summary cards:", sub_items=[
        "Total registered users (Service Seekers + Service Providers)",
        "Total verified Service Providers",
        "Total service requests submitted",
        "Total completed services",
        "Total platform revenue (accumulated from 5% commission on all transactions)",
        "Number of active users on the platform",
    ])
    add_numbered_step(doc, 3, "Statistics update in real-time as users interact with the platform, providing an accurate snapshot of platform activity.")

    add_body_paragraph(doc,
        "These statistics are essential for tracking the platform's growth, identifying trends, "
        "and making data-driven decisions about platform improvements and marketing strategies."
    )

    # ---- 4.3.4.6 Handling Support Messages ----
    add_heading_text(doc, "4.3.4.6 Handling Support Messages", "subsubsection")

    add_body_paragraph(doc,
        "The admin is responsible for handling support inquiries from both Service Seekers and "
        "Service Providers through the platform's built-in support chat."
    )

    add_numbered_step(doc, 1, 'Click "Support" in the admin sidebar navigation.')
    add_numbered_step(doc, 2, "View the list of active support chat conversations from users.")
    add_numbered_step(doc, 3, "Click on a conversation to open the chat and view the user's messages.")
    add_numbered_step(doc, 4, "Type your response in the message field and click send.")
    add_numbered_step(doc, 5, "The user receives your response in real-time.")
    add_numbered_step(doc, 6, "Continue the conversation until the user's issue is resolved.")
    add_numbered_step(doc, 7, "Mark the conversation as resolved when appropriate.")

    add_body_paragraph(doc,
        "Timely and effective support handling is crucial for maintaining user satisfaction and "
        "retention on the platform. The admin should aim to respond to support inquiries promptly "
        "and provide clear, helpful guidance."
    )

    add_page_break(doc)

    # ========================================================================
    # 4.4 Sample System Run
    # ========================================================================
    add_heading_text(doc, "4.4 Sample System Run", "section")

    add_body_paragraph(doc,
        "This section presents a narrative walkthrough of a complete end-to-end transaction on "
        "the HomeEase platform. The scenario demonstrates how all the components of the system "
        "work together seamlessly, from initial registration through to payment completion and "
        "review. This walkthrough follows two fictional users: Ada (a Service Seeker) and Emeka "
        "(a Service Provider)."
    )

    add_body_paragraph(doc, "The Complete HomeEase Transaction Flow:", bold=True)

    add_numbered_step(doc, 1, "Ada, a homeowner in Lagos, opens her web browser and navigates to https://homeease.vercel.app. The HomeEase landing page loads, displaying the platform's services and features.")
    add_numbered_step(doc, 2, "Ada clicks \"Get Started\" and selects \"Service Seeker\" as her role. She fills out the registration form with her name, email address, phone number, and creates a password. She clicks \"Sign Up\" and her account is created successfully.")
    add_numbered_step(doc, 3, "Ada is redirected to the login page. She enters her email and password, and clicks \"Sign In.\" She is taken to her personalised Service Seeker dashboard, where she can see an overview of her account activity.")
    add_numbered_step(doc, 4, "Ada needs a plumber to fix a leaking kitchen pipe. She clicks \"Find Artisans\" in the sidebar navigation. The artisan search page opens, displaying a grid of verified service providers.")
    add_numbered_step(doc, 5, "Ada types \"Plumber\" in the search bar and presses Enter. A filtered list of plumbers appears. She browses through the results and notices \"Emeka Plumbing Services\" with an impressive 4.8-star rating based on 24 reviews. His skills include Plumbing, Pipe Fitting, and Bathroom Installation, with an hourly rate of ₦3,500.")
    add_numbered_step(doc, 6, "Ada clicks \"Book Now\" on Emeka's profile card. The booking form appears. She fills in the details:", sub_items=[
        "Location: 14 Adeola Street, Ikeja, Lagos",
        "Preferred Date: 15th March, 2025",
        "Preferred Time: 10:00 AM",
        'Description: "Kitchen pipe under the sink is leaking. Need it fixed urgently."',
    ])
    add_numbered_step(doc, 7, "Ada clicks \"Confirm Booking.\" A success message appears: \"Your booking request has been sent to Emeka Plumbing Services.\" The request status is set to PENDING.")
    add_numbered_step(doc, 8, "Meanwhile, Emeka is logged into his provider dashboard. He hears a notification sound and sees a new job offer in the \"Job Offers\" section. He clicks on it and sees Ada's booking details: plumbing service at her Ikeja address on 15th March at 10:00 AM.")
    add_numbered_step(doc, 9, "Emeka reviews the details and clicks \"Accept.\" The job moves to his \"My Jobs\" section with a status of ACCEPTED. A notification is automatically sent to Ada: \"Emeka Plumbing Services has accepted your booking request.\"")
    add_numbered_step(doc, 10, "Ada receives the notification and sees that her request status has updated to ACCEPTED. She can now message Emeka directly through the platform to discuss any additional details about the plumbing issue.")
    add_numbered_step(doc, 11, "On the morning of 15th March, Emeka arrives at Ada's residence at 9:55 AM. He opens the HomeEase app on his phone, navigates to \"My Jobs,\" and clicks \"Check In\" on Ada's job. A live timer begins counting from 00:00:00.")
    add_numbered_step(doc, 12, "The LIVE indicator confirms the timer is running. Ada can see the live timer on her end as well from the \"My Requests\" section, providing full transparency of the service duration.")
    add_numbered_step(doc, 13, "Emeka assesses the leaking pipe, replaces a worn-out section, and seals all connections. After 2 hours and 15 minutes of work, the repair is complete and the pipe is functioning properly.")
    add_numbered_step(doc, 14, "Emeka clicks \"Check Out\" on the job. The timer stops at 02:15:00 (2 hours 15 minutes). The job status changes to COMPLETED. Ada receives a notification: \"Your plumbing service by Emeka Plumbing Services has been completed.\"")
    add_numbered_step(doc, 15, "Ada reviews the completed work and is satisfied. She goes to \"My Requests,\" clicks on the completed job, and clicks \"Pay Now.\" She enters the amount of ₦5,000 (calculated as ₦3,500 per hour × 2.25 hours = ₦7,875, but Ada and Emeka had agreed on a fixed price of ₦5,000).")
    add_numbered_step(doc, 16, 'Ada clicks "Proceed to Payment" and is redirected to the Paystack secure payment page. She enters her debit card details and authorises the payment of ₦5,000. The payment is processed successfully.')
    add_numbered_step(doc, 17, "Ada is redirected back to HomeEase. The payment status updates to PAID. The platform automatically calculates the commission: 5% of ₦5,000 equals ₦250. Emeka's net payout is ₦4,750 (₦5,000 − ₦250).")
    add_numbered_step(doc, 18, "A rating prompt appears on Ada's completed request. She clicks 5 stars and writes: \"Excellent work! Emeka was punctual, professional, and fixed the leak perfectly. Highly recommended.\" She clicks \"Submit Review.\"")
    add_numbered_step(doc, 19, "Emeka's average rating updates to reflect Ada's 5-star review. The review is visible on his public profile, helping future clients make informed decisions.")
    add_numbered_step(doc, 20, "Emeka receives notifications confirming the payment and the positive review. The ₦4,750 is added to his available wallet balance. Emeka can now initiate a withdrawal request to transfer the funds to his bank account.")
    add_numbered_step(doc, 21, "Emeka navigates to the \"Earnings\" section and clicks \"Withdraw.\" He enters ₦4,750 as the withdrawal amount and confirms. The request is sent to the admin for processing.")
    add_numbered_step(doc, 22, "The admin receives the payout request, reviews the details, and clicks \"Initiate Payout.\" The system transfers ₦4,750 to Emeka's registered bank account via Paystack Transfer. Emeka receives a notification confirming the successful payout.")

    add_body_paragraph(doc,
        "This complete walkthrough demonstrates how HomeEase seamlessly connects Service Seekers "
        "with Service Providers, manages the entire service lifecycle from booking to payment, "
        "provides real-time tracking and communication, handles secure payment processing with "
        "automatic commission calculation, and maintains a transparent review system. Every step "
        "of the process is handled within the platform, eliminating the need for external "
        "communication or manual payment arrangements."
    )

    add_page_break(doc)

    # ========================================================================
    # 4.5 Summary
    # ========================================================================
    add_heading_text(doc, "4.5 Summary", "section")

    add_body_paragraph(doc,
        "This chapter has provided a comprehensive overview of the HomeEase system implementation "
        "environment and a detailed, step-by-step user guide for every process on the platform. "
        "The implementation environment section documented the hardware and software tools utilised "
        "during the development and deployment of the application, including Next.js 16, React 19, "
        "TypeScript 5, Tailwind CSS 4, shadcn/ui, Prisma ORM, PostgreSQL (via Supabase), Socket.IO, "
        "NextAuth.js v4, Paystack API, and Vercel for deployment."
    )

    add_body_paragraph(doc,
        "The user manual section provided exhaustive walkthroughs for all three user roles on the "
        "platform. For Service Seekers, the guide covered accessing the website, creating an account, "
        "navigating the dashboard, finding and booking service providers, viewing requests, making "
        "payments via Paystack, rating providers, using the real-time messaging system, managing "
        "notifications, and editing profile information."
    )

    add_body_paragraph(doc,
        "For Service Providers, the manual detailed the verification process, dashboard navigation, "
        "managing incoming job offers, the check-in and check-out workflow with live time tracking, "
        "viewing transaction histories, managing earnings and wallet withdrawals, messaging clients, "
        "and maintaining professional profiles."
    )

    add_body_paragraph(doc,
        "For Administrators, the guide covered accessing the admin dashboard, verifying new service "
        "provider applications, managing all registered users, processing payout requests via "
        "Paystack Transfer, monitoring platform statistics, and handling support inquiries."
    )

    add_body_paragraph(doc,
        "Finally, the sample system run section presented a narrative end-to-end scenario that "
        "demonstrated how all platform components work together in a real-world transaction, from "
        "initial user registration through to service completion, payment processing, review "
        "submission, and fund withdrawal. This walkthrough confirms that the HomeEase platform "
        "successfully fulfils its objective of connecting homeowners with verified artisans "
        "through a streamlined, transparent, and efficient digital platform."
    )

    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    doc.save(OUTPUT_PATH)
    print(f"Document saved successfully to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


# ==============================================================================
# MAIN
# ==============================================================================

if __name__ == "__main__":
    generate_chapter4()
