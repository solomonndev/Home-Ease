"""
Script to update HomeEase project writeup:
1. Update tech stack to match actual implementation
2. Add inline academic references
"""

from docx import Document
import copy

INPUT_FILE = "/home/z/my-project/upload/new mariam.docx"
OUTPUT_FILE = "/home/z/my-project/upload/HomeEase_Project_Writeup.docx"

doc = Document(INPUT_FILE)

# ============================================================
# TECH STACK REPLACEMENTS
# ============================================================

tech_stack_replacements = {
    # --- Section 1.E Methodology ---
    "HTML5, CSS3, and JavaScript,for responsive and dynamic user interfaces":
        "Next.js 16 (a React-based framework) with TypeScript for type-safe, responsive, and dynamic user interfaces",

    "The backend system is implemented using Node.js with Express which manage business logic, service scheduling, user authentication, and communication between system components.":
        "The backend system is implemented using Next.js API Routes (App Router) with the Bun JavaScript runtime, which manages business logic, service scheduling, user authentication, real-time messaging via Socket.io, and communication between system components.",

    "Data is stored and managed using a relational database such as MySQL for scalability and data structure requirements.":
        "Data is stored and managed using SQLite as the relational database engine, with Prisma ORM providing a type-safe schema definition language, automated migration management, and an intuitive query API for data access.",

    "Firebase is used for secure user authentication and role-based access control, while Stripe or Paystack APIs are employed for secure payment processing.":
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication and role-based access control, while the Paystack payment gateway API is employed for secure payment processing integrated with a virtual wallet and escrow system.",

    "Cloud services such as Google Cloud, or Hostinger support hosting, storage, and scalability of the virtual space.":
        "Vercel, a cloud platform optimized for Next.js applications, supports hosting, serverless function execution, and automatic scaling of the virtual space.",

    "software testing tools such as Jest, Selenium, or Postman":
        "software testing tools such as Jest for unit testing, and Postman for API endpoint verification",

    # --- Section 3.4.1 Software Requirements Table ---
    "React native + \n                                                 Tailwind CSS":
        "Next.js 16 + TypeScript +\n                                                 Tailwind CSS 4 + shadcn/ui",

    "React native + Tailwind CSS":
        "Next.js 16 + TypeScript +\n Tailwind CSS 4",

    "Node.js with\n        Express/API routes":
        "Next.js 16 (App Router) +\n        Bun Runtime + Socket.io",

    "Node.js with Express/API routes":
        "Next.js 16 (App Router) + Bun Runtime + Socket.io",

    "MongoDB(NOSQL)":
        "SQLite with Prisma ORM",

    "MongoDB (NOSQL)":
        "SQLite with Prisma ORM",

    "Paystack or\n        Flutterwave":
        "Paystack Payment\n        Gateway",

    "Paystack or Flutterwave":
        "Paystack Payment Gateway",

    # --- Section 3.4 intro paragraph ---
    "Commonly used DBMS solutions such as MySQL or MongoDB are suitable due to their reliability and scalability.":
        "SQLite, chosen for its serverless architecture, zero-configuration deployment, and reliability, managed through Prisma ORM which provides type-safe database access and schema management.",
}

def replace_in_paragraphs(doc, replacements):
    """Replace text in all paragraphs, handling run splitting."""
    count = 0
    for para in doc.paragraphs:
        full_text = para.text
        modified = False
        
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                # We need to do the replacement
                # First, save paragraph formatting
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        count += 1
                        modified = True
        
        # For multi-run matches (text split across runs)
        if not modified:
            for old_text, new_text in replacements.items():
                if old_text in full_text:
                    # Clear all runs and set combined text on first run
                    combined = full_text.replace(old_text, new_text)
                    if combined != full_text:
                        # Clear all run texts
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = combined
                        else:
                            para.add_run(combined)
                        count += 1
                        modified = True
    
    return count

def replace_in_tables(doc, replacements):
    """Replace text in all table cells."""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    for old_text, new_text in replacements.items():
                        if old_text in full_text:
                            combined = full_text.replace(old_text, new_text)
                            if combined != full_text:
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = combined
                                count += 1
    return count

print("Applying tech stack replacements to paragraphs...")
p_count = replace_in_paragraphs(doc, tech_stack_replacements)
print(f"  Paragraph replacements: {p_count}")

print("Applying tech stack replacements to tables...")
t_count = replace_in_tables(doc, tech_stack_replacements)
print(f"  Table replacements: {t_count}")

# ============================================================
# INLINE REFERENCES ADDITIONS
# ============================================================

# Map of key phrases to inline citations to add
# Format: (search_text, replacement_text_with_citation)
inline_references = [
    # Chapter 1 - Introduction/Background
    (
        "The rapid advancement of digital technologies has transformed how services are delivered across various industries.",
        "The rapid advancement of digital technologies has transformed how services are delivered across various industries [1]."
    ),
    (
        "Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping.",
        "Domestic services remain essential to households across the globe, particularly in areas such as childcare, elder care, and housekeeping [2]."
    ),
    (
        "Domestic services refer to a range of household-related activities carried out to support the daily functioning, comfort, and well-being of individuals or families within a home.",
        "Domestic services refer to a range of household-related activities carried out to support the daily functioning, comfort, and well-being of individuals or families within a home [3]."
    ),
    (
        "The concept of a virtual space for domestic service involves creating an online platform or digital environment where domestic service providers and users can interact seamlessly.",
        "The concept of a virtual space for domestic service involves creating an online platform or digital environment where domestic service providers and users can interact seamlessly [4]."
    ),
    
    # Chapter 1 - Literature Review Table references
    (
        "Khatri & Gupta (2020)",
        "Khatri and Gupta [5]"
    ),
    (
        "Adeyemi & Fatile",
        "Adeyemi and Fatile [6]"
    ),
    (
        "Rana et al. (2019)",
        "Rana et al. [7]"
    ),
    (
        "Indravasan et al. (2018)",
        "Indravasan et al. [8]"
    ),
    (
        "Chen et al. (2021)",
        "Chen et al. [9]"
    ),

    # Chapter 1 - Problem Statement
    (
        "The provision of domestic services is largely managed through informal and manual processes",
        "The provision of domestic services is largely managed through informal and manual processes [5, 6]"
    ),

    # Chapter 2 - Theoretical Framework
    (
        "The Technology Acceptance Model (TAM) posits that users\u2019 acceptance and use of a technological system are determined by perceived usefulness and perceived ease of use (Davis, 1989).",
        "The Technology Acceptance Model (TAM) posits that users\u2019 acceptance and use of a technological system are determined by perceived usefulness and perceived ease of use [10]."
    ),
    (
        "The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries",
        "The Platform Economy Theory provides a foundation for understanding virtual spaces as digital intermediaries [11]"
    ),
    (
        "Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms.",
        "Additionally, Service Quality Theory explains how users evaluate services delivered through digital platforms [12]."
    ),

    # Chapter 2 - Virtual Space
    (
        "Core functional modules of the platform include service discovery and matching, booking and scheduling",
        "Core functional modules of the platform include service discovery and matching, booking and scheduling [13]"
    ),
    (
        "When integrated with Paystack digital payment system",
        "When integrated with a digital payment system such as Paystack [14]"
    ),
    (
        "Digital payment systems have achieved widespread adoption due to their high transaction speed",
        "Digital payment systems have achieved widespread adoption due to their high transaction speed [14]"
    ),

    # Chapter 2 - Web 3.0
    (
        "Web 3.0 refers to the next stage in the evolution of the internet",
        "Web 3.0 refers to the next stage in the evolution of the internet [15]"
    ),
    (
        "blockchain is a distributed ledger maintained by a network of computers",
        "blockchain is a distributed ledger maintained by a network of computers [16]"
    ),

    # Chapter 2 - Database
    (
        "A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated.",
        "A database is an organized collection of data that is stored electronically and designed to be easily accessed, managed, and updated [17]."
    ),
    (
        "Examples of widely used DBMSs include MySQL, PostgreSQL, Oracle Database, Microsoft SQL Server, and MongoDB.",
        "Examples of widely used DBMSs include MySQL, PostgreSQL, Oracle Database, Microsoft SQL Server, and MongoDB [18]."
    ),
    (
        "Relational databases, organize data into tables consisting of rows and columns.",
        "Relational databases organize data into tables consisting of rows and columns [19]."
    ),
    (
        "ACID properties are fundamental principles that ensure the reliability and correctness of database transactions.",
        "ACID properties are fundamental principles that ensure the reliability and correctness of database transactions [20]."
    ),
    (
        "Prisma ORM providing a type-safe schema definition language",
        "Prisma ORM [21] providing a type-safe schema definition language"
    ),

    # Chapter 2 - DBMS
    (
        "A Database Management System (DBMS) is software that allows users to create, store, manage, and retrieve data efficiently and securely.",
        "A Database Management System (DBMS) is software that allows users to create, store, manage, and retrieve data efficiently and securely [18]."
    ),
    (
        "SQL is an example of a database access language",
        "SQL is an example of a database access language [22]"
    ),

    # Chapter 2 - Data Models
    (
        "A data model refers to an abstract representation of data structures that are used to organize and manage data in a database or information system.",
        "A data model refers to an abstract representation of data structures that are used to organize and manage data in a database or information system [23]."
    ),

    # Chapter 2 - Data Modeling
    (
        "Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system.",
        "Data modeling refers to the process of creating a visual and logical representation of data requirements and structures within an information system [24]."
    ),
    (
        "Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures.",
        "Entity\u2013Relationship (ER) Modeling: ER modeling uses entities, attributes, and relationships to visually represent data structures [25]."
    ),
    (
        "Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency.",
        "Normalization: Normalization is the process of organizing data to eliminate redundancy and dependency [26]."
    ),

    # Chapter 2 - Software Process Models / SDLC
    (
        "A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems.",
        "A software process model is a structured framework that defines the sequence of activities involved in the development, deployment, and maintenance of software systems [27]."
    ),
    (
        "System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner.",
        "System Development Life Cycle (SDLC) is a systematic framework used to develop information systems in an organized and efficient manner [28]."
    ),
    (
        "The waterfall model is a sequential, plan driven-process",
        "The waterfall model is a sequential, plan-driven process [29]"
    ),
    (
        "The V model (Verification and Validation model) is an extension of the waterfall model.",
        "The V model (Verification and Validation model) is an extension of the waterfall model [30]."
    ),
    (
        "The agile process model encourages continuous iterations of development and testing.",
        "The agile process model encourages continuous iterations of development and testing [31]."
    ),

    # Chapter 2 - Software Hosting
    (
        "Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure.",
        "Software hosting is a fundamental concept in modern software engineering that refers to the deployment, execution, and management of software applications on remote computing infrastructure [32]."
    ),
    (
        "Platform as a Service (PaaS) provides a complete hosting environment",
        "Platform as a Service (PaaS) provides a complete hosting environment [33]"
    ),

    # Chapter 2 - Related Works
    (
        "Aishwaryalakshmi et al. (2024),",
        "Aishwaryalakshmi et al. [34],"
    ),
    (
        "Pais and Zanoni (2024),",
        "Pais and Zanoni [35],"
    ),
    (
        "Meyanban et al. (2024),",
        "Meyanban et al. [36],"
    ),
    (
        "Orth and Baum (2024)",
        "Orth and Baum [37]"
    ),
    (
        "Rakhewar et al. (2023):",
        "Rakhewar et al. [38]:"
    ),
    (
        "Yadav et al. (2023)",
        "Yadav et al. [39]"
    ),
    (
        "Sehgal & Yathrath (2022):",
        "Sehgal and Yathrath [40]:"
    ),
    (
        "Adeyemi and Fatile (2021),",
        "Adeyemi and Fatile [6],"
    ),
    (
        "Khatri and Gupta (2020 ),",
        "Khatri and Gupta [5],"
    ),
    (
        "Chatterjee et al. (2021)",
        "Chatterjee et al. [41]"
    ),
    (
        "Vallas and Schor (2020)",
        "Vallas and Schor [42]"
    ),
    (
        "Rana et al. (2019):",
        "Rana et al. [7]:"
    ),
    (
        "Indravasan et al. (2018)",
        "Indravasan et al. [8]"
    ),
    (
        "Berg, E. Rani (2018),",
        "Berg [43],"
    ),
    (
        "Sundararajan (2016)",
        "Sundararajan [44]"
    ),

    # Chapter 3 - Methodology
    (
        "The system implementation follows an agile development methodology",
        "The system implementation follows an agile development methodology [31]"
    ),
    (
        "Next.js 16 (a React-based framework) with TypeScript",
        "Next.js 16 (a React-based framework) with TypeScript [45]"
    ),
    (
        "Next.js API Routes (App Router) with the Bun JavaScript runtime",
        "Next.js API Routes (App Router) with the Bun JavaScript runtime [45, 46]"
    ),
    (
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication",
        "JSON Web Tokens (JWT) are used for secure, stateless user authentication [47]"
    ),
    (
        "Paystack payment gateway API is employed for secure payment processing",
        "Paystack payment gateway API is employed for secure payment processing [14]"
    ),
    (
        "Vercel, a cloud platform optimized for Next.js applications",
        "Vercel, a cloud platform optimized for Next.js applications [48]"
    ),
    
    # Chapter 3 - Requirements / Design
    (
        "Intelligent Matching and Scheduling",
        "Intelligent Matching and Scheduling [49]"
    ),
    (
        "A Use Case Diagram illustrates the primary interactions between the system actors and the system itself.",
        "A Use Case Diagram illustrates the primary interactions between the system actors and the system itself [50]."
    ),
    (
        "A Context Diagram (also known as a Level 0 Data Flow Diagram)",
        "A Context Diagram (also known as a Level 0 Data Flow Diagram) [51]"
    ),
    (
        "A Data Flow Diagram (DFD) is a graphical and logical representation",
        "A Data Flow Diagram (DFD) is a graphical and logical representation [52]"
    ),
    (
        "An Entity Relationship Diagram (ERD) is a conceptual blueprint of your database.",
        "An Entity Relationship Diagram (ERD) is a conceptual blueprint of the database [25]."
    ),
    (
        "The Relational Data Model (RDM), also known as the Physical Database Schema",
        "The Relational Data Model (RDM), also known as the Physical Database Schema [53]"
    ),
]

print("\nApplying inline references...")
ref_count = 0
for para in doc.paragraphs:
    full_text = para.text
    modified = False
    
    for search_text, replacement_text in inline_references:
        if search_text in full_text and replacement_text not in full_text:
            combined = full_text.replace(search_text, replacement_text)
            if combined != full_text:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = combined
                else:
                    para.add_run(combined)
                ref_count += 1
                modified = True
                full_text = combined
    
# Also apply references in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full_text = para.text
                for search_text, replacement_text in inline_references:
                    if search_text in full_text and replacement_text not in full_text:
                        combined = full_text.replace(search_text, replacement_text)
                        if combined != full_text:
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = combined
                            else:
                                para.add_run(combined)
                            ref_count += 1

print(f"  Inline reference additions: {ref_count}")

# ============================================================
# SAVE OUTPUT
# ============================================================
doc.save(OUTPUT_FILE)
print(f"\nDocument saved to: {OUTPUT_FILE}")
print("Done!")
